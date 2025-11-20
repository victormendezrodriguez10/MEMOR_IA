# memoria_tecnica_pro_mejorada.py
"""
MEMOR.IA - Plataforma Profesional de Generación de Memorias Técnicas
Sistema completo con login y extracción automática
Versión 5.0 - MEJORADA con enfoque en criterios de valoración
"""

import streamlit as st
import openai
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime, timedelta
import io
import json
import PyPDF2
import re
import base64
import requests
from bs4 import BeautifulSoup  # Para extraer datos de web
import urllib.parse
import time
import hashlib
import os
from dotenv import load_dotenv
import sqlite3
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import secrets
import string
from datetime import datetime, timedelta

# Cargar variables de entorno
load_dotenv()

# ============ CONFIGURACIÓN ============
# Usar variable de entorno para mayor seguridad
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)
MODELO_IA = "claude-opus-4-1-20250805"

# Configuración de email
EMAIL_HOST = os.getenv("EMAIL_HOST", "oclem-com.correoseguro.dinaserver.com")
EMAIL_PORT = os.getenv("EMAIL_PORT", 465)
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

# Configuración de Stripe
STRIPE_PUBLIC_KEY = os.getenv("STRIPE_PUBLIC_KEY")
STRIPE_SECRET_KEY = os.getenv("STRIPE_SECRET_KEY")

# Precio del servicio
PRECIO_SERVICIO = 363  # 363€ IVA incluido

# ============ SISTEMA DE IDIOMAS ============
IDIOMAS = {
    "es": {
        "nombre": "Español",
        "bandera": "🇪🇸",
        "login_title": "🧠 MEMOR.IA",
        "login_subtitle": "Generador Inteligente de Memorias Técnicas",
        "login_form": "Iniciar Sesión",
        "email": "Email",
        "password": "Contraseña",
        "login_button": "Iniciar Sesión",
        "demo_button": "Demo Rápida",
        "register_button": "Nuevo Cliente - Registrarse",
        "forgot_password": "¿Olvidaste tu contraseña?",
        "welcome": "¡Bienvenido",
        "register_title": "Registro de Nuevo Cliente",
        "full_name": "Nombre Completo",
        "company": "Empresa/Razón Social",
        "phone": "Teléfono",
        "address": "Dirección Completa",
        "price_info": "**Precio:** 363€ IVA incluido",
        "bank_account": "Número de Cuenta Bancaria (IBAN)",
        "accept_terms": "He leído y acepto los términos y condiciones del contrato de servicios",
        "register_pay_button": "🚀 REGISTRARSE Y PAGAR (363€)",
        "admin_panel": "👑 PANEL DE ADMINISTRACIÓN",
        "user_management": "Gestión de Usuarios y Pagos",
        "memory_generator": "Generador Inteligente de Memorias Técnicas"
    },
    "en": {
        "nombre": "English",
        "bandera": "🇬🇧",
        "login_title": "🧠 MEMOR.IA",
        "login_subtitle": "Intelligent Technical Reports Generator",
        "login_form": "Sign In",
        "email": "Email",
        "password": "Password",
        "login_button": "Sign In",
        "demo_button": "Quick Demo",
        "register_button": "New Client - Register",
        "forgot_password": "Forgot your password?",
        "welcome": "Welcome",
        "register_title": "New Client Registration",
        "full_name": "Full Name",
        "company": "Company Name",
        "phone": "Phone",
        "address": "Full Address",
        "price_info": "**Price:** 363€ VAT included",
        "bank_account": "Bank Account Number (IBAN)",
        "accept_terms": "I have read and accept the terms and conditions of the service contract",
        "register_pay_button": "🚀 REGISTER AND PAY (363€)",
        "admin_panel": "👑 ADMIN PANEL",
        "user_management": "User and Payment Management",
        "memory_generator": "Intelligent Technical Reports Generator"
    },
    "de": {
        "nombre": "Deutsch",
        "bandera": "🇩🇪",
        "login_title": "🧠 MEMOR.IA",
        "login_subtitle": "Intelligenter Generator für technische Berichte",
        "login_form": "Anmelden",
        "email": "E-Mail",
        "password": "Passwort",
        "login_button": "Anmelden",
        "demo_button": "Schnelle Demo",
        "register_button": "Neuer Kunde - Registrieren",
        "forgot_password": "Passwort vergessen?",
        "welcome": "Willkommen",
        "register_title": "Registrierung neuer Kunden",
        "full_name": "Vollständiger Name",
        "company": "Firmenname",
        "phone": "Telefon",
        "address": "Vollständige Adresse",
        "price_info": "**Preis:** 363€ inkl. MwSt.",
        "bank_account": "Bankkontonummer (IBAN)",
        "accept_terms": "Ich habe die Allgemeinen Geschäftsbedingungen gelesen und akzeptiert",
        "register_pay_button": "🚀 REGISTRIEREN UND BEZAHLEN (363€)",
        "admin_panel": "👑 ADMIN-PANEL",
        "user_management": "Benutzer- und Zahlungsverwaltung",
        "memory_generator": "Intelligenter Generator für technische Berichte"
    },
    "pt": {
        "nombre": "Português",
        "bandera": "🇵🇹",
        "login_title": "🧠 MEMOR.IA",
        "login_subtitle": "Gerador Inteligente de Relatórios Técnicos",
        "login_form": "Iniciar Sessão",
        "email": "Email",
        "password": "Palavra-passe",
        "login_button": "Iniciar Sessão",
        "demo_button": "Demo Rápida",
        "register_button": "Novo Cliente - Registar",
        "forgot_password": "Esqueceu a palavra-passe?",
        "welcome": "Bem-vindo",
        "register_title": "Registo de Novo Cliente",
        "full_name": "Nome Completo",
        "company": "Nome da Empresa",
        "phone": "Telefone",
        "address": "Endereço Completo",
        "price_info": "**Preço:** 363€ IVA incluído",
        "bank_account": "Número de Conta Bancária (IBAN)",
        "accept_terms": "Li e aceito os termos e condições do contrato de serviços",
        "register_pay_button": "🚀 REGISTAR E PAGAR (363€)",
        "admin_panel": "👑 PAINEL DE ADMINISTRAÇÃO",
        "user_management": "Gestão de Utilizadores e Pagamentos",
        "memory_generator": "Gerador Inteligente de Relatórios Técnicos"
    },
    "fr": {
        "nombre": "Français",
        "bandera": "🇫🇷",
        "login_title": "🧠 MEMOR.IA",
        "login_subtitle": "Générateur Intelligent de Rapports Techniques",
        "login_form": "Se connecter",
        "email": "Email",
        "password": "Mot de passe",
        "login_button": "Se connecter",
        "demo_button": "Démo Rapide",
        "register_button": "Nouveau Client - S'inscrire",
        "forgot_password": "Mot de passe oublié?",
        "welcome": "Bienvenue",
        "register_title": "Inscription Nouveau Client",
        "full_name": "Nom Complet",
        "company": "Nom de l'Entreprise",
        "phone": "Téléphone",
        "address": "Adresse Complète",
        "price_info": "**Prix:** 363€ TVA incluse",
        "bank_account": "Numéro de Compte Bancaire (IBAN)",
        "accept_terms": "J'ai lu et j'accepte les termes et conditions du contrat de services",
        "register_pay_button": "🚀 S'INSCRIRE ET PAYER (363€)",
        "admin_panel": "👑 PANNEAU D'ADMINISTRATION",
        "user_management": "Gestion des Utilisateurs et Paiements",
        "memory_generator": "Générateur Intelligent de Rapports Techniques"
    }
}

# Configuración de la página - DEBE ir al inicio
st.set_page_config(
    page_title="MEMOR.IA - Generador Inteligente de Memorias Técnicas",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Configuración para subpath (comentado para evitar errores DOM en local)
# import os
# if os.getenv('STREAMLIT_SERVER_BASE_URL_PATH'):
#     st.markdown("""
#     <script>
#     // Ajustar rutas para subpath
#     window.location.pathname = window.location.pathname.replace('/memor-ia/memor-ia', '/memor-ia');
#     </script>
#     """, unsafe_allow_html=True)

def get_text(key):
    """Obtiene el texto en el idioma seleccionado"""
    idioma = st.session_state.get('idioma', 'es')
    return IDIOMAS[idioma].get(key, IDIOMAS['es'][key])

def validar_iban(iban):
    """
    Valida un número IBAN según el estándar internacional
    Retorna (es_valido, mensaje_error)
    """
    if not iban:
        return False, "El número de cuenta es obligatorio"

    # Limpiar el IBAN (quitar espacios y convertir a mayúsculas)
    iban_limpio = iban.replace(' ', '').replace('-', '').upper()

    # Verificar longitud mínima y máxima
    if len(iban_limpio) < 15 or len(iban_limpio) > 34:
        return False, "El IBAN debe tener entre 15 y 34 caracteres"

    # Verificar que empiece con dos letras
    if not iban_limpio[:2].isalpha():
        return False, "El IBAN debe empezar con el código de país (2 letras)"

    # Verificar que los siguientes 2 caracteres sean dígitos
    if not iban_limpio[2:4].isdigit():
        return False, "Los dígitos de control deben ser números"

    # Verificar que el resto sean alfanuméricos
    if not iban_limpio[4:].isalnum():
        return False, "El IBAN contiene caracteres no válidos"

    # Algoritmo de validación IBAN (módulo 97)
    try:
        # Mover los primeros 4 caracteres al final
        iban_reordenado = iban_limpio[4:] + iban_limpio[:4]

        # Convertir letras a números (A=10, B=11, ..., Z=35)
        iban_numerico = ""
        for caracter in iban_reordenado:
            if caracter.isdigit():
                iban_numerico += caracter
            else:
                iban_numerico += str(ord(caracter) - ord('A') + 10)

        # Verificar módulo 97
        if int(iban_numerico) % 97 == 1:
            return True, ""
        else:
            return False, "El IBAN no pasa la verificación de dígitos de control"

    except (ValueError, OverflowError):
        return False, "Error en la validación del IBAN"

# ============ SISTEMA DE AUTENTICACIÓN Y BASE DE DATOS ============

def init_database():
    """Inicializa la base de datos SQLite"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    # Tabla de usuarios
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            nombre TEXT NOT NULL,
            empresa TEXT NOT NULL,
            telefono TEXT,
            cif TEXT,
            direccion TEXT,
            numero_cuenta TEXT,
            rol TEXT DEFAULT 'Usuario',
            fecha_registro DATETIME DEFAULT CURRENT_TIMESTAMP,
            activo BOOLEAN DEFAULT 1,
            plan TEXT DEFAULT 'basico',
            fecha_expiracion DATE
        )
    ''')

    # Tabla de perfiles de empresa (datos que se reutilizan en memorias)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS perfiles_empresa (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER UNIQUE,
            sector TEXT,
            empleados TEXT,
            experiencia_anos TEXT,
            certificaciones TEXT,  -- JSON con array de certificaciones
            otras_certificaciones TEXT,
            experiencia_similar TEXT,
            logo_path TEXT,
            medios_materiales TEXT,
            herramientas_software TEXT,
            equipo_tecnico TEXT,  -- JSON con array de técnicos
            fecha_actualizacion DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (usuario_id) REFERENCES usuarios (id)
        )
    ''')

    # Tabla de pagos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS pagos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            stripe_payment_id TEXT,
            importe REAL,
            fecha_pago DATETIME DEFAULT CURRENT_TIMESTAMP,
            estado TEXT DEFAULT 'pendiente',
            plan TEXT,
            FOREIGN KEY (usuario_id) REFERENCES usuarios (id)
        )
    ''')

    # Tabla de tokens de recuperación
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tokens_recuperacion (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT NOT NULL,
            token TEXT NOT NULL,
            fecha_creacion DATETIME DEFAULT CURRENT_TIMESTAMP,
            usado BOOLEAN DEFAULT 0
        )
    ''')

    # Usuario administrador por defecto
    cursor.execute('''
        INSERT OR IGNORE INTO usuarios
        (email, password, nombre, empresa, rol, activo)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', ("vmendez@oclem.com", hashlib.sha256("favorito1998".encode()).hexdigest(),
          "Víctor Méndez", "OCLEM", "Administrador", 1))

    # Usuario demo
    cursor.execute('''
        INSERT OR IGNORE INTO usuarios
        (email, password, nombre, empresa, rol, activo)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', ("demo@demo.com", hashlib.sha256("demo123".encode()).hexdigest(),
          "Usuario Demo", "Empresa Demo", "Usuario", 1))

    conn.commit()
    conn.close()

def generar_password():
    """Genera una contraseña aleatoria"""
    caracteres = string.ascii_letters + string.digits
    return ''.join(secrets.choice(caracteres) for i in range(12))

def enviar_email(destinatario, asunto, mensaje_html):
    """Envía un email"""
    try:
        # Debug: verificar configuración
        print(f"EMAIL_HOST: {EMAIL_HOST}")
        print(f"EMAIL_PORT: {EMAIL_PORT}")
        print(f"EMAIL_USER: {EMAIL_USER}")
        print(f"Destinatario: {destinatario}")

        msg = MIMEMultipart('alternative')
        msg['Subject'] = asunto
        msg['From'] = EMAIL_USER
        msg['To'] = destinatario

        html_part = MIMEText(mensaje_html, 'html')
        msg.attach(html_part)

        # Puerto 465 requiere SSL directo, no STARTTLS
        if int(EMAIL_PORT) == 465:
            server = smtplib.SMTP_SSL(EMAIL_HOST, int(EMAIL_PORT))
        else:
            # Puerto 587 usa STARTTLS
            server = smtplib.SMTP(EMAIL_HOST, int(EMAIL_PORT))
            server.starttls()

        # Configurar timeout para evitar que se cuelgue
        server.timeout = 10
        server.login(EMAIL_USER, EMAIL_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        print(f"Error enviando email: {e}")
        st.error(f"Error detallado: {str(e)}")
        return False

def verificar_credenciales(email, password):
    """Verifica las credenciales del usuario"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    password_hash = hashlib.sha256(password.encode()).hexdigest()

    # Debug: verificar si el usuario existe
    cursor.execute('SELECT email, activo FROM usuarios WHERE email = ?', (email,))
    usuario_existe = cursor.fetchone()

    if not usuario_existe:
        print(f"DEBUG: Usuario {email} no encontrado en la base de datos")
        conn.close()
        return False, None

    if not usuario_existe[1]:
        print(f"DEBUG: Usuario {email} existe pero está inactivo")
        conn.close()
        return False, None

    cursor.execute('''
        SELECT nombre, empresa, rol, activo FROM usuarios
        WHERE email = ? AND password = ?
    ''', (email, password_hash))

    resultado = cursor.fetchone()
    conn.close()

    if resultado and resultado[3]:  # activo = 1
        print(f"DEBUG: Login exitoso para {email}")
        return True, {
            "nombre": resultado[0],
            "empresa": resultado[1],
            "rol": resultado[2]
        }

    print(f"DEBUG: Credenciales incorrectas para {email}")
    return False, None

def registrar_usuario(datos_usuario):
    """Registra un nuevo usuario en la base de datos"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    password = generar_password()
    password_hash = hashlib.sha256(password.encode()).hexdigest()

    try:
        cursor.execute('''
            INSERT INTO usuarios
            (email, password, nombre, empresa, telefono, cif, direccion, numero_cuenta, plan, fecha_expiracion)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (datos_usuario['email'], password_hash, datos_usuario['nombre'],
              datos_usuario['empresa'], datos_usuario['telefono'], datos_usuario['cif'],
              datos_usuario['direccion'], datos_usuario.get('numero_cuenta', ''), 'basico',
              (datetime.now() + timedelta(days=30)).date()))

        conn.commit()
        conn.close()
        return password
    except sqlite3.IntegrityError:
        conn.close()
        return None

# La base de datos se inicializa en main()

def obtener_perfil_empresa(usuario_email):
    """Obtiene el perfil de empresa de un usuario"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    cursor.execute('''
        SELECT pe.* FROM perfiles_empresa pe
        JOIN usuarios u ON pe.usuario_id = u.id
        WHERE u.email = ?
    ''', (usuario_email,))

    resultado = cursor.fetchone()
    conn.close()

    if resultado:
        return {
            'sector': resultado[2],
            'empleados': resultado[3],
            'experiencia_anos': resultado[4],
            'certificaciones': json.loads(resultado[5]) if resultado[5] else [],
            'otras_certificaciones': resultado[6],
            'experiencia_similar': resultado[7],
            'logo_path': resultado[8],
            'medios_materiales': resultado[9],
            'herramientas_software': resultado[10],
            'equipo_tecnico': json.loads(resultado[11]) if resultado[11] else [],
            'fecha_actualizacion': resultado[12]
        }
    return None

def guardar_perfil_empresa(usuario_email, datos_perfil):
    """Guarda o actualiza el perfil de empresa de un usuario"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    # Obtener ID del usuario
    cursor.execute('SELECT id FROM usuarios WHERE email = ?', (usuario_email,))
    usuario_result = cursor.fetchone()

    if not usuario_result:
        conn.close()
        return False

    usuario_id = usuario_result[0]

    # Verificar si ya existe un perfil
    cursor.execute('SELECT id FROM perfiles_empresa WHERE usuario_id = ?', (usuario_id,))
    existe = cursor.fetchone()

    # Preparar datos JSON
    certificaciones_json = json.dumps(datos_perfil.get('certificaciones', []))
    equipo_tecnico_json = json.dumps(datos_perfil.get('equipo_tecnico', []))

    if existe:
        # Actualizar perfil existente
        cursor.execute('''
            UPDATE perfiles_empresa SET
                sector = ?, empleados = ?, experiencia_anos = ?,
                certificaciones = ?, otras_certificaciones = ?,
                experiencia_similar = ?, logo_path = ?,
                medios_materiales = ?, herramientas_software = ?,
                equipo_tecnico = ?, fecha_actualizacion = CURRENT_TIMESTAMP
            WHERE usuario_id = ?
        ''', (
            datos_perfil.get('sector', ''),
            datos_perfil.get('empleados', ''),
            datos_perfil.get('experiencia_anos', ''),
            certificaciones_json,
            datos_perfil.get('otras_certificaciones', ''),
            datos_perfil.get('experiencia_similar', ''),
            datos_perfil.get('logo_path', ''),
            datos_perfil.get('medios_materiales', ''),
            datos_perfil.get('herramientas_software', ''),
            equipo_tecnico_json,
            usuario_id
        ))
    else:
        # Crear nuevo perfil
        cursor.execute('''
            INSERT INTO perfiles_empresa
            (usuario_id, sector, empleados, experiencia_anos,
             certificaciones, otras_certificaciones, experiencia_similar,
             logo_path, medios_materiales, herramientas_software, equipo_tecnico)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            usuario_id,
            datos_perfil.get('sector', ''),
            datos_perfil.get('empleados', ''),
            datos_perfil.get('experiencia_anos', ''),
            certificaciones_json,
            datos_perfil.get('otras_certificaciones', ''),
            datos_perfil.get('experiencia_similar', ''),
            datos_perfil.get('logo_path', ''),
            datos_perfil.get('medios_materiales', ''),
            datos_perfil.get('herramientas_software', ''),
            equipo_tecnico_json
        ))

    conn.commit()
    conn.close()
    return True

def guardar_logo_usuario(usuario_email, logo_file):
    """Guarda el logo de un usuario específico"""
    import os

    # Crear directorio para logos si no existe (usar directorio actual en lugar de hardcoded)
    logos_dir = "/Users/macintosh/Desktop/memoria copia/logos_usuarios"
    if not os.path.exists(logos_dir):
        os.makedirs(logos_dir)

    # Generar nombre único para el logo
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    usuario_safe = usuario_email.replace("@", "_").replace(".", "_")
    extension = logo_file.name.split('.')[-1]
    logo_filename = f"logo_{usuario_safe}_{timestamp}.{extension}"
    logo_path = os.path.join(logos_dir, logo_filename)

    # Guardar archivo
    try:
        with open(logo_path, "wb") as f:
            f.write(logo_file.getbuffer())
        print(f"Logo guardado en: {logo_path}")  # Debug
        return logo_path
    except Exception as e:
        print(f"Error guardando logo: {e}")  # Debug
        return None

def guardar_documentos_anexos(usuario_email, archivos_subidos, categoria):
    """Guarda los documentos anexos de un usuario específico"""
    import os
    from datetime import datetime

    if not archivos_subidos:
        return []

    # Crear directorio para documentos si no existe
    docs_dir = "/Users/macintosh/Desktop/memoria copia/documentos_usuarios"
    usuario_dir = os.path.join(docs_dir, usuario_email.replace("@", "_").replace(".", "_"))

    if not os.path.exists(usuario_dir):
        os.makedirs(usuario_dir)

    documentos_guardados = []
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    for i, archivo in enumerate(archivos_subidos):
        # Generar nombre único para el documento
        extension = archivo.name.split('.')[-1]
        nombre_base = archivo.name.replace(f".{extension}", "")
        doc_filename = f"{nombre_base}_{timestamp}_{i}.{extension}"
        doc_path = os.path.join(usuario_dir, doc_filename)

        # Guardar archivo
        try:
            with open(doc_path, "wb") as f:
                f.write(archivo.getbuffer())

            # Añadir información del documento
            documentos_guardados.append({
                'nombre': archivo.name,
                'categoria': categoria,
                'ruta_archivo': doc_path,
                'fecha_subida': datetime.now().strftime("%d/%m/%Y %H:%M"),
                'tamaño': len(archivo.getbuffer())
            })
        except Exception as e:
            st.error(f"Error guardando {archivo.name}: {str(e)}")

    return documentos_guardados

def generar_token_recuperacion(email):
    """Genera un token para recuperación de contraseña"""
    token = secrets.token_urlsafe(32)

    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    cursor.execute('''
        INSERT INTO tokens_recuperacion (email, token)
        VALUES (?, ?)
    ''', (email, token))

    conn.commit()
    conn.close()
    return token

def validar_token_recuperacion(email, token):
    """Valida un token de recuperación"""
    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    cursor.execute('''
        SELECT id FROM tokens_recuperacion
        WHERE email = ? AND token = ? AND usado = 0
        AND datetime(fecha_creacion) > datetime('now', '-1 hour')
    ''', (email, token))

    resultado = cursor.fetchone()
    conn.close()

    return resultado is not None

def cambiar_password(email, nueva_password, token):
    """Cambia la contraseña usando un token válido"""
    if not validar_token_recuperacion(email, token):
        return False

    conn = sqlite3.connect('memoria_usuarios.db')
    cursor = conn.cursor()

    password_hash = hashlib.sha256(nueva_password.encode()).hexdigest()

    # Actualizar contraseña
    cursor.execute('''
        UPDATE usuarios SET password = ? WHERE email = ?
    ''', (password_hash, email))

    # Marcar token como usado
    cursor.execute('''
        UPDATE tokens_recuperacion SET usado = 1 WHERE email = ? AND token = ?
    ''', (email, token))

    conn.commit()
    conn.close()
    return True

def mostrar_recuperacion():
    """Página de recuperación de contraseña"""
    aplicar_estilos_login()


    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        st.markdown("""
        <div class="login-container">
            <div class="logo-container">
                <h1 class="app-title">🧠 MEMOR.IA</h1>
                <p class="app-subtitle">Recuperar Contraseña</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

        if 'recovery_step' not in st.session_state:
            st.session_state.recovery_step = 1

        if st.session_state.recovery_step == 1:
            # Paso 1: Solicitar email
            st.markdown("### 📧 Introduce tu email")

            with st.form("recovery_form_step1"):
                email_recovery = st.text_input("Email", placeholder="tu@email.com")
                enviar_token = st.form_submit_button("Enviar código de recuperación", use_container_width=True)

                if enviar_token and email_recovery:
                    # Verificar que el email existe
                    conn = sqlite3.connect('memoria_usuarios.db')
                    cursor = conn.cursor()
                    cursor.execute("SELECT id FROM usuarios WHERE email = ?", (email_recovery,))
                    existe = cursor.fetchone()
                    conn.close()

                    if existe:
                        token = generar_token_recuperacion(email_recovery)

                        mensaje_html = f"""
                        <html>
                        <body>
                            <h2>Recuperación de contraseña - MEMOR.IA</h2>
                            <p>Has solicitado recuperar tu contraseña.</p>

                            <h3>Código de recuperación:</h3>
                            <p style="font-size: 24px; font-weight: bold; background: #f0f0f0; padding: 10px; border-radius: 5px;">
                                {token[:8].upper()}
                            </p>

                            <p><em>Este código expira en 1 hora.</em></p>

                            <p>Si no solicitaste este cambio, ignora este email.</p>
                            <p><em>Equipo MEMOR.IA</em></p>
                        </body>
                        </html>
                        """

                        if enviar_email(email_recovery, "Recuperación de contraseña - MEMOR.IA", mensaje_html):
                            st.success("✅ Código enviado a tu email")
                            st.session_state.recovery_email = email_recovery
                            st.session_state.recovery_step = 2
                            st.rerun()
                        else:
                            st.error("❌ Error enviando el email")
                    else:
                        st.error("❌ Email no encontrado")

        elif st.session_state.recovery_step == 2:
            # Paso 2: Validar código y nueva contraseña
            st.markdown("### 🔑 Introduce el código y nueva contraseña")
            st.info(f"📧 Código enviado a: {st.session_state.recovery_email}")

            with st.form("recovery_form_step2"):
                codigo = st.text_input("Código de recuperación", placeholder="ABC12345")
                nueva_password = st.text_input("Nueva contraseña", type="password")
                confirmar_password = st.text_input("Confirmar contraseña", type="password")

                cambiar = st.form_submit_button("Cambiar contraseña", use_container_width=True)

                if cambiar:
                    if not codigo or not nueva_password or not confirmar_password:
                        st.error("❌ Completa todos los campos")
                    elif nueva_password != confirmar_password:
                        st.error("❌ Las contraseñas no coinciden")
                    elif len(nueva_password) < 6:
                        st.error("❌ La contraseña debe tener al menos 6 caracteres")
                    else:
                        # Buscar token completo
                        conn = sqlite3.connect('memoria_usuarios.db')
                        cursor = conn.cursor()
                        cursor.execute('''
                            SELECT token FROM tokens_recuperacion
                            WHERE email = ? AND token LIKE ? AND usado = 0
                            AND datetime(fecha_creacion) > datetime('now', '-1 hour')
                        ''', (st.session_state.recovery_email, codigo.upper() + '%'))
                        token_completo = cursor.fetchone()
                        conn.close()

                        if token_completo and cambiar_password(st.session_state.recovery_email, nueva_password, token_completo[0]):
                            st.success("✅ Contraseña cambiada correctamente")
                            st.balloons()
                            st.session_state.recovery_step = 1
                            del st.session_state.recovery_email
                            time.sleep(2)
                            st.session_state.mostrar_login = True
                            st.session_state.mostrar_recuperacion = False
                            st.rerun()
                        else:
                            st.error("❌ Código inválido o expirado")

        # Botón volver
        if st.button("← Volver al Login", use_container_width=True):
            st.session_state.mostrar_recuperacion = False
            st.session_state.mostrar_login = True
            st.session_state.recovery_step = 1
            if 'recovery_email' in st.session_state:
                del st.session_state.recovery_email
            st.rerun()

# ============ ESTILOS CSS PARA LOGIN ============
def aplicar_estilos_login():
    st.markdown("""
    <style>
        /* Fondo animado */
        .stApp {
            background: linear-gradient(-45deg, #667eea, #764ba2, #f093fb, #4facfe);
            background-size: 400% 400%;
            animation: gradientAnimation 15s ease infinite;
        }
        
        @keyframes gradientAnimation {
            0% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
            100% { background-position: 0% 50%; }
        }
        
        /* Contenedor de login */
        .login-container {
            background: rgba(255, 255, 255, 0.95);
            backdrop-filter: blur(10px);
            border-radius: 20px;
            padding: 3rem;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            max-width: 450px;
            margin: auto;
            animation: slideDown 0.5s ease-out;
        }
        
        @keyframes slideDown {
            from {
                opacity: 0;
                transform: translateY(-50px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        /* Logo y título */
        .logo-container {
            text-align: center;
            margin-bottom: 2rem;
            animation: pulse 2s infinite;
        }
        
        @keyframes pulse {
            0% { transform: scale(1); }
            50% { transform: scale(1.05); }
            100% { transform: scale(1); }
        }
        
        .app-title {
            font-size: 3.5rem;
            font-weight: 900;
            color: #2c3e50;
            margin: 0;
            letter-spacing: -2px;
            text-shadow: 2px 2px 4px rgba(255,255,255,0.8);
        }

        .app-subtitle {
            color: #34495e;
            font-size: 1.2rem;
            margin-top: 0.5rem;
            font-weight: 500;
            text-shadow: 1px 1px 2px rgba(255,255,255,0.8);
        }
        
        /* Campos de entrada */
        .stTextInput > div > div > input {
            background: white;
            border: 2px solid #e0e0e0;
            border-radius: 10px;
            padding: 12px 15px;
            font-size: 16px;
            transition: all 0.3s ease;
        }
        
        .stTextInput > div > div > input:focus {
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }
        
        /* Features */
        .features-container {
            display: flex;
            justify-content: space-around;
            margin-top: 3rem;
            flex-wrap: wrap;
        }
        
        .feature-card {
            background: rgba(255, 255, 255, 0.9);
            padding: 1.5rem;
            border-radius: 15px;
            text-align: center;
            width: 200px;
            margin: 10px;
            transition: all 0.3s ease;
            box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        }
        
        .feature-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        
        .feature-icon {
            font-size: 2.5rem;
            margin-bottom: 0.5rem;
        }
    </style>
    """, unsafe_allow_html=True)

# ============ ESTILOS CSS PARA LA APLICACIÓN ============
def aplicar_estilos_app():
    st.markdown("""
    <style>
        .main-header {
            text-align: center;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2rem;
            border-radius: 15px;
            margin-bottom: 2rem;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        
        .logo-memoria {
            font-size: 2.5rem;
            font-weight: 900;
            color: white;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }
        
        .section-header {
            color: #1f4e79;
            border-bottom: 3px solid #667eea;
            padding-bottom: 0.5rem;
            margin: 2rem 0 1rem;
        }
        
        .info-box {
            background: #e8f4f8;
            padding: 1.5rem;
            border-radius: 10px;
            border-left: 5px solid #667eea;
            margin: 1rem 0;
        }
        
        .success-box {
            background: #d4edda;
            padding: 1rem;
            border-radius: 8px;
            border-left: 5px solid #28a745;
            margin: 1rem 0;
        }
        
        .warning-box {
            background: #fff3cd;
            border-left: 5px solid #ffc107;
            padding: 1.5rem;
            border-radius: 10px;
            margin: 2rem 0;
        }
        
        .stButton>button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            font-weight: bold;
            border: none;
            padding: 0.75rem 2rem;
            border-radius: 10px;
            width: 100%;
            transition: all 0.3s ease;
        }
        
        .stButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 15px rgba(102, 126, 234, 0.4);
        }

        /* Centrar logos */
        .stImage {
            display: flex !important;
            justify-content: center !important;
            align-items: center !important;
        }

        [data-testid="stImage"] {
            display: flex !important;
            justify-content: center !important;
            align-items: center !important;
        }

    </style>
    """, unsafe_allow_html=True)

# ============ FUNCIONES DE GENERACIÓN DE ENCABEZADO ============
def agregar_encabezado_con_logo(doc, logo_path, datos_empresa):
    """
    Añade encabezado con logo de empresa a todas las páginas
    """
    try:
        print(f"DEBUG: Intentando añadir logo desde: {logo_path}")  # Debug
        print(f"DEBUG: Logo existe: {os.path.exists(logo_path) if logo_path else False}")  # Debug

        # Acceder al encabezado de la primera sección
        section = doc.sections[0]
        header = section.header

        # Crear tabla en el encabezado para logo y texto
        header_table = header.add_table(rows=1, cols=2, width=Inches(6))
        header_table.autofit = False

        # Celda izquierda: Logo
        left_cell = header_table.rows[0].cells[0]
        left_cell.width = Inches(2)

        # Intentar añadir logo con más verificaciones
        logo_añadido = False
        if logo_path:
            print(f"DEBUG: Verificando logo en ruta: {logo_path}")
            if os.path.exists(logo_path):
                try:
                    paragraph = left_cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(1.5))
                    logo_añadido = True
                    print("DEBUG: Logo añadido correctamente")
                except Exception as e:
                    print(f"DEBUG: Error añadiendo imagen: {e}")
            else:
                print(f"DEBUG: Archivo de logo no encontrado en: {logo_path}")

        # Si no se pudo añadir logo, añadir texto placeholder
        if not logo_añadido:
            paragraph = left_cell.paragraphs[0]
            run = paragraph.add_run("🏢 LOGO")
            run.font.size = Pt(12)
            run.font.bold = True

        # Celda derecha: Información de empresa
        right_cell = header_table.rows[0].cells[1]
        right_cell.width = Inches(4)
        paragraph = right_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run(f"{datos_empresa.get('razon_social', 'EMPRESA')}\n")
        run.font.size = Pt(10)
        run.font.bold = True

        run = paragraph.add_run(f"CIF: {datos_empresa.get('cif', '')}\n")
        run.font.size = Pt(9)

        run = paragraph.add_run(f"{datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(9)

        # Línea separadora
        header.add_paragraph('_' * 50)

        print("DEBUG: Encabezado añadido correctamente")

    except Exception as e:
        print(f"DEBUG: Error en agregar_encabezado_con_logo: {e}")
        # Añadir un encabezado básico en caso de error
        try:
            section = doc.sections[0]
            header = section.header
            p = header.paragraphs[0]
            p.text = f"{datos_empresa.get('razon_social', 'EMPRESA')} - {datetime.now().strftime('%d/%m/%Y')}"
        except:
            pass

# ============ FORMULARIO DE REGISTRO Y PAGO ============
def mostrar_registro():
    """Página de registro con pasarela de pago"""
    aplicar_estilos_login()

    # Selector de idioma en la esquina superior
    col_lang, col_space = st.columns([1, 4])
    with col_lang:
        idioma_seleccionado = st.selectbox(
            "🌍",
            options=list(IDIOMAS.keys()),
            format_func=lambda x: f"{IDIOMAS[x]['bandera']} {IDIOMAS[x]['nombre']}",
            index=list(IDIOMAS.keys()).index(st.session_state.get('idioma', 'es')),
            key="selector_idioma_registro"
        )
        st.session_state.idioma = idioma_seleccionado


    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        # Logo centrado en div
        try:
            # Convertir imagen a base64 para incluir en HTML
            import base64
            with open("logo.png", "rb") as img_file:
                logo_base64 = base64.b64encode(img_file.read()).decode()

            st.markdown(f"""
            <div style="text-align: center; margin-bottom: 1rem; background: transparent; padding: 1rem;">
                <img src="data:image/png;base64,{logo_base64}" width="200" style="margin-bottom: 0.5rem;">
                <h2 style="color: #2c3e50; margin-top: 0.5rem; font-weight: 600;">{get_text('register_title')}</h2>
            </div>
            """, unsafe_allow_html=True)
        except:
            st.markdown(f"""
            <div style="text-align: center; margin-bottom: 1rem; background: transparent; padding: 1rem;">
                <h1 class="app-title">{get_text('login_title')}</h1>
                <h2 style="color: #2c3e50; margin-top: 0.5rem; font-weight: 600;">{get_text('register_title')}</h2>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("### 📝 Datos de Registro")

        with st.form("registro_form"):
            col_a, col_b = st.columns(2)

            with col_a:
                nombre = st.text_input(f"{get_text('full_name')} *", placeholder="Juan Pérez")
                empresa = st.text_input(f"{get_text('company')} *", placeholder="Empresa S.L.")
                email = st.text_input(f"{get_text('email')} *", placeholder="tu@email.com")

            with col_b:
                telefono = st.text_input(f"{get_text('phone')} *", placeholder="+34 600 123 456")
                cif = st.text_input("CIF/NIF *", placeholder="B12345678")
                direccion = st.text_area(f"{get_text('address')} *", placeholder="Calle, número, ciudad...")

            # Plan de servicio
            st.markdown("### 💰 Plan Seleccionado")
            st.info("**Plan Básico - 363€ IVA incluido**\n- Memorias técnicas ilimitadas\n- Soporte técnico\n- Actualizaciones incluidas")

            # Checkbox de aceptación
            acepta_contrato = st.checkbox(
                f"{get_text('accept_terms')} *",
                help="Obligatorio para proceder con el registro"
            )

            # Información de pago
            st.markdown("### 💳 Datos Bancarios")
            st.info(f"💰 {get_text('price_info')}\n📧 Recibirás las credenciales por email tras el registro\n📞 Te contactaremos para gestionar el pago")

            numero_cuenta = st.text_input(
                get_text('bank_account'),
                placeholder="ES21 1465 0100 72 2030876293",
                help="Formato IBAN completo para domiciliación bancaria"
            )

            # Validación en tiempo real del IBAN
            if numero_cuenta:
                es_valido, mensaje_error = validar_iban(numero_cuenta)
                if es_valido:
                    st.success("✅ IBAN válido")
                else:
                    st.error(f"❌ {mensaje_error}")

            # Botón de registro y pago
            registrar = st.form_submit_button(get_text('register_pay_button'), use_container_width=True)

            if registrar:
                # Validaciones
                if not all([nombre, empresa, email, telefono, cif, direccion]):
                    st.error("❌ Por favor, completa todos los campos obligatorios")
                elif not acepta_contrato:
                    st.error("❌ Debes aceptar los términos y condiciones")
                elif not numero_cuenta:
                    st.error("❌ Por favor, introduce el número de cuenta bancaria")
                else:
                    # Validar el IBAN antes de procesar
                    es_valido_iban, mensaje_error_iban = validar_iban(numero_cuenta)
                    if not es_valido_iban:
                        st.error(f"❌ Error en el número de cuenta: {mensaje_error_iban}")
                    else:
                        # Simular procesamiento de pago
                        with st.spinner("Procesando pago y creando cuenta..."):
                            time.sleep(2)  # Simular procesamiento

                            datos_usuario = {
                                'nombre': nombre,
                                'empresa': empresa,
                                'email': email,
                                'telefono': telefono,
                                'cif': cif,
                                'direccion': direccion,
                                'numero_cuenta': numero_cuenta
                            }

                            password = registrar_usuario(datos_usuario)

                            if password:
                                # Enviar email con credenciales
                                mensaje_html = f"""
                                <html>
                                <body>
                                    <h2>¡Bienvenido a MEMOR.IA!</h2>
                                    <p>Tu cuenta ha sido creada exitosamente.</p>

                                    <h3>Datos de acceso:</h3>
                                    <p><strong>Email:</strong> {email}</p>
                                    <p><strong>Contraseña:</strong> {password}</p>

                                    <p>Ya puedes acceder a: <a href="http://localhost:8502">MEMOR.IA</a></p>

                                    <p>¡Gracias por confiar en nosotros!</p>
                                    <p><em>Equipo MEMOR.IA</em></p>
                                </body>
                                </html>
                                """

                                if enviar_email(email, "Bienvenido a MEMOR.IA - Datos de acceso", mensaje_html):
                                    st.success("✅ ¡Registro completado! Revisa tu email para acceder.")
                                    st.balloons()
                                    time.sleep(3)
                                    st.session_state.mostrar_login = True
                                    st.rerun()
                                else:
                                    st.warning("⚠️ Cuenta creada pero error enviando email. Contacta con soporte.")
                            else:
                                st.error("❌ El email ya está registrado")

        # Enlaces
        if st.button("← Volver al Login", use_container_width=True):
            st.session_state.mostrar_registro = False
            st.session_state.mostrar_login = True
            st.rerun()

# ============ PÁGINA DE LOGIN ============
def mostrar_login():
    aplicar_estilos_login()

    # Selector de idioma en la esquina superior
    col_lang, col_space = st.columns([1, 4])
    with col_lang:
        if 'idioma' not in st.session_state:
            st.session_state.idioma = 'es'

        idioma_seleccionado = st.selectbox(
            "🌍",
            options=list(IDIOMAS.keys()),
            format_func=lambda x: f"{IDIOMAS[x]['bandera']} {IDIOMAS[x]['nombre']}",
            index=list(IDIOMAS.keys()).index(st.session_state.idioma),
            key="selector_idioma_login"
        )
        st.session_state.idioma = idioma_seleccionado


    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        # Logo centrado en div
        try:
            # Convertir imagen a base64 para incluir en HTML
            import base64
            with open("logo.png", "rb") as img_file:
                logo_base64 = base64.b64encode(img_file.read()).decode()

            st.markdown(f"""
            <div style="text-align: center; margin-bottom: 1rem; background: transparent; padding: 1rem;">
                <img src="data:image/png;base64,{logo_base64}" width="200" style="margin-bottom: 0.5rem;">
                <h3 style="color: #2c3e50; margin-top: 0.5rem; font-weight: 500;">{get_text('login_subtitle')}</h3>
            </div>
            """, unsafe_allow_html=True)
        except:
            st.markdown(f"""
            <div style="text-align: center; margin-bottom: 1rem; background: transparent; padding: 1rem;">
                <h1 class="app-title">{get_text('login_title')}</h1>
                <h3 style="color: #2c3e50; margin-top: 0.5rem; font-weight: 500;">{get_text('login_subtitle')}</h3>
            </div>
            """, unsafe_allow_html=True)
        
        with st.form("login_form"):
            st.markdown(f"### 🔐 {get_text('login_form')}")

            email = st.text_input(f"📧 {get_text('email')}", placeholder="tu@email.com")
            password = st.text_input(f"🔑 {get_text('password')}", type="password", placeholder="••••••••")

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                submit = st.form_submit_button(get_text('login_button'), use_container_width=True)
            with col_btn2:
                demo = st.form_submit_button(get_text('demo_button'), use_container_width=True)
            
            if submit:
                if email and password:
                    valido, datos_usuario = verificar_credenciales(email, password)
                    if valido:
                        st.session_state.logged_in = True
                        st.session_state.user_data = datos_usuario
                        st.session_state.user_email = email
                        st.success("✅ ¡Bienvenido " + datos_usuario["nombre"] + "!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("❌ Credenciales incorrectas")
                else:
                    st.warning("⚠️ Por favor, completa todos los campos")
            
            if demo:
                # Login automático como demo
                valido, datos_usuario = verificar_credenciales("demo@demo.com", "demo123")
                if valido:
                    st.session_state.logged_in = True
                    st.session_state.user_data = datos_usuario
                    st.session_state.user_email = "demo@demo.com"
                    st.session_state.demo_mode = True  # Modo demo solo visualización
                    st.success("🎭 Entrando en modo demostración...")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ Error en el modo demo")

        # Opciones adicionales
        st.markdown("---")

        col_reg, col_rec = st.columns(2)

        with col_reg:
            if st.button(f"👤 {get_text('register_button')}", use_container_width=True):
                st.session_state.mostrar_registro = True
                st.session_state.mostrar_login = False
                st.rerun()

        with col_rec:
            if st.button(f"🔑 {get_text('forgot_password')}", use_container_width=True):
                st.session_state.mostrar_recuperacion = True
                st.session_state.mostrar_login = False
                st.rerun()
    
    # Features
    st.markdown("""
    <div class="features-container">
        <div class="feature-card">
            <div class="feature-icon">🤖</div>
            <div class="feature-title">IA Avanzada</div>
            <div class="feature-desc">para resultados profesionales</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">📊</div>
            <div class="feature-title">Criterios de Valoración</div>
            <div class="feature-desc">Desarrollo detallado de cada criterio</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">🏢</div>
            <div class="feature-title">Entrenada con más de 30 años de experiencia </div>
            <div class="feature-desc">elaborando memorias técnicas</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">⚡</div>
            <div class="feature-title">Rápido</div>
            <div class="feature-desc">Memorias en minutos</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ============ FUNCIONES MEJORADAS DE GENERACIÓN ============

def llamar_ia_mejorado(prompt, max_tokens=3000, temperature=0.3):
    """Versión mejorada usando Anthropic Claude"""
    try:
        # Debug: verificar configuración
        print(f"DEBUG IA: API Key configurada: {bool(ANTHROPIC_API_KEY)}")
        print(f"DEBUG IA: Modelo: {MODELO_IA}")
        print(f"DEBUG IA: Longitud del prompt: {len(prompt)}")

        if not ANTHROPIC_API_KEY:
            st.error("❌ API Key de Anthropic no configurada. Verifica tu archivo .env")
            return "Error: API Key no encontrada. Contacta al administrador."

        # Usar Claude Opus 4.1 - la última versión más potente
        modelo_a_usar = MODELO_IA  # claude-opus-4-1-20250805

        # Claude puede manejar prompts mucho más largos
        if len(prompt) > 50000:
            # Dividir en partes si es muy largo
            partes = []
            texto = prompt
            while len(texto) > 50000:
                corte = texto[:50000].rfind('\n')
                if corte == -1:
                    corte = 50000
                partes.append(texto[:corte])
                texto = texto[corte:]
            if texto:
                partes.append(texto)

            respuesta_completa = ""
            for i, parte in enumerate(partes):
                print(f"DEBUG IA: Procesando parte {i+1}/{len(partes)}")
                response = anthropic_client.messages.create(
                    model=modelo_a_usar,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    system="Eres un experto redactor de memorias técnicas profesionales para licitaciones. Redacta siempre en párrafos largos y fluidos, sin usar listas, viñetas, asteriscos ni guiones.",
                    messages=[
                        {"role": "user", "content": parte}
                    ]
                )
                respuesta_completa += response.content[0].text + "\n\n"
            return respuesta_completa
        else:
            print("DEBUG IA: Enviando prompt a Claude...")
            response = anthropic_client.messages.create(
                model=modelo_a_usar,
                max_tokens=max_tokens,
                temperature=temperature,
                system="Eres un experto redactor de memorias técnicas profesionales para licitaciones. Redacta siempre en párrafos largos y fluidos, sin usar listas, viñetas, asteriscos ni guiones.",
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            print("DEBUG IA: Respuesta recibida exitosamente")
            return response.content[0].text
    except Exception as e:
        error_msg = f"Error con Claude: {str(e)}"
        print(f"DEBUG IA ERROR: {error_msg}")
        st.error(error_msg)
        return f"Error en la generación: {str(e)}"

# ============ FUNCIONES AUXILIARES PARA SECTORES ============

def detectar_sector_proyecto(objeto, texto_ppt):
    """Detecta el sector del proyecto basado en palabras clave"""
    texto_completo = f"{objeto} {texto_ppt}".lower()

    sectores = {
        'construccion': ['construcción', 'obra', 'edificación', 'estructura', 'cimentación', 'hormigón', 'edificio', 'reforma'],
        'electricidad': ['eléctrica', 'instalación eléctrica', 'suministro eléctrico', 'transformador', 'cableado', 'iluminación', 'energía'],
        'software': ['software', 'aplicación', 'sistema informático', 'desarrollo', 'programación', 'base de datos', 'servidor', 'web'],
        'mantenimiento': ['mantenimiento', 'conservación', 'reparación', 'limpieza', 'jardinería', 'servicios'],
        'consultoría': ['consultoría', 'asesoramiento', 'estudio', 'análisis', 'auditoría', 'evaluación'],
        'suministros': ['suministro', 'material', 'equipamiento', 'mobiliario', 'compra', 'adquisición'],
        'transporte': ['transporte', 'logística', 'vehículo', 'distribución', 'traslado']
    }

    for sector, palabras_clave in sectores.items():
        if any(palabra in texto_completo for palabra in palabras_clave):
            return sector

    return 'general'

def get_rol_profesional(sector):
    """Devuelve el rol profesional correspondiente al sector"""
    roles = {
        'construccion': 'Ingeniero de Caminos, Canales y Puertos especializado en construcción',
        'electricidad': 'Ingeniero Eléctrico especializado en instalaciones',
        'software': 'Ingeniero en Informática especializado en desarrollo de software',
        'mantenimiento': 'Técnico especialista en mantenimiento industrial',
        'consultoría': 'Consultor senior especializado en gestión empresarial',
        'suministros': 'Especialista en aprovisionamiento y logística',
        'transporte': 'Especialista en logística y transporte',
        'general': 'Profesional especializado en licitaciones públicas'
    }
    return roles.get(sector, roles['general'])

def generar_estructura_especifica(nombre_criterio, sector):
    """Genera estructura específica según el criterio y sector"""
    criterio_lower = nombre_criterio.lower()

    # Estructuras específicas por tipo de criterio
    if 'personal' in criterio_lower or 'equipo' in criterio_lower or 'técnico' in criterio_lower:
        return """
        1. ANÁLISIS DE COMPETENCIAS REQUERIDAS
        Evaluación técnica de las competencias específicas necesarias para el proyecto,
        identificando perfiles profesionales clave y niveles de especialización requeridos.

        2. SELECCIÓN Y ASIGNACIÓN DE PERSONAL
        Proceso de selección del equipo técnico, criterios de asignación basados en experiencia
        específica y metodología de distribución de responsabilidades técnicas.

        3. ORGANIZACIÓN TÉCNICA DEL EQUIPO
        Estructura organizativa del equipo de trabajo, definición de roles específicos,
        cadena de supervisión técnica y protocolos de coordinación.

        4. CUALIFICACIÓN Y CERTIFICACIONES
        Acreditaciones profesionales del personal asignado, formación específica continuada
        y certificaciones técnicas especializadas vigentes.

        5. DISPONIBILIDAD Y DEDICACIÓN
        Planificación de disponibilidad temporal del personal, dedicación exclusiva o parcial,
        y mecanismos de sustitución ante contingencias.
        """

    elif 'metodolog' in criterio_lower or 'enfoque' in criterio_lower:
        return f"""
        1. FUNDAMENTACIÓN TÉCNICA DE LA METODOLOGÍA
        Análisis técnico de la metodología propuesta basada en normativas específicas del sector {sector},
        justificación de la elección metodológica y adaptación a las características del proyecto.

        2. DESARROLLO OPERATIVO DE PROCESOS
        Descripción detallada de los procedimientos operativos, secuencias de ejecución técnica,
        protocolos específicos de actuación y sistemática de trabajo.

        3. HERRAMIENTAS Y TECNOLOGÍAS APLICADAS
        Especificación técnica de herramientas especializadas, software específico del sector,
        tecnologías aplicadas y sistemas de medición y control.

        4. CONTROL DE CALIDAD INTEGRADO
        Sistema de aseguramiento de la calidad técnica, puntos de control críticos,
        indicadores de rendimiento y protocolos de verificación.

        5. OPTIMIZACIÓN Y MEJORA CONTINUA
        Mecanismos de optimización de procesos, sistemas de retroalimentación técnica,
        y procedimientos de mejora continua basados en métricas específicas.
        """

    elif 'experiencia' in criterio_lower or 'capacidad' in criterio_lower:
        return f"""
        1. TRAYECTORIA ESPECIALIZADA EN EL SECTOR
        Análisis de la experiencia específica acumulada en proyectos similares del sector {sector},
        evolución de competencias técnicas y especialización progresiva.

        2. PROYECTOS DE REFERENCIA RELEVANTES
        Descripción técnica de proyectos ejecutados de características similares,
        análisis de complejidad técnica abordada y resultados obtenidos.

        3. CAPACIDADES TÉCNICAS DESARROLLADAS
        Competencias técnicas específicas desarrolladas a través de la experiencia práctica,
        conocimiento de normativa especializada y dominio de tecnologías específicas.

        4. APRENDIZAJE Y ADAPTACIÓN TÉCNICA
        Capacidad de adaptación a nuevas tecnologías del sector, actualización técnica continuada
        y incorporación de innovaciones en metodologías de trabajo.

        5. RECONOCIMIENTO PROFESIONAL Y CERTIFICACIONES
        Acreditaciones profesionales obtenidas, reconocimientos sectoriales
        y participación en organismos técnicos especializados.
        """

    elif 'recursos' in criterio_lower or 'medios' in criterio_lower:
        return f"""
        1. INVENTARIO TÉCNICO DE RECURSOS
        Catalogación detallada de recursos técnicos disponibles, especificaciones técnicas completas
        de equipamiento y evaluación de capacidades operativas.

        2. INFRAESTRUCTURA TÉCNICA ESPECIALIZADA
        Descripción de instalaciones técnicas específicas, laboratorios especializados,
        talleres técnicos y espacios de trabajo especializados para el sector {sector}.

        3. TECNOLOGÍA Y EQUIPAMIENTO AVANZADO
        Inventario de tecnología especializada disponible, especificaciones técnicas de equipos,
        software especializado y herramientas de última generación.

        4. SISTEMAS DE GESTIÓN DE RECURSOS
        Metodología de gestión y asignación de recursos, sistemas de control de disponibilidad,
        protocolos de mantenimiento preventivo y gestión de inventarios técnicos.

        5. CAPACIDAD DE ESCALADO Y ADAPTACIÓN
        Flexibilidad en la asignación de recursos según fases del proyecto,
        capacidad de incorporación de recursos adicionales y adaptación a requerimientos específicos.
        """

    elif 'planificac' in criterio_lower or 'control' in criterio_lower:
        return f"""
        1. ARQUITECTURA DE PLANIFICACIÓN TÉCNICA
        Metodología de planificación especializada para proyectos del sector {sector},
        definición de fases técnicas críticas y secuenciación de actividades especializadas.

        2. CRONOGRAMA TÉCNICO DETALLADO
        Desarrollo de cronograma técnico específico, identificación de rutas críticas,
        planificación de recursos especializados y coordinación interdisciplinar.

        3. SISTEMA DE MONITORIZACIÓN Y CONTROL
        Implementación de sistemas de seguimiento técnico en tiempo real,
        indicadores clave de rendimiento específicos y protocolos de control de calidad.

        4. GESTIÓN DE RIESGOS TÉCNICOS
        Identificación de riesgos técnicos específicos del sector, planes de contingencia técnica,
        estrategias de mitigación y protocolos de respuesta ante incidencias.

        5. REPORTING Y COMUNICACIÓN TÉCNICA
        Sistema de informes técnicos periódicos, comunicación con stakeholders especializados
        y documentación técnica de seguimiento y control.
        """

    else:
        # Estructura genérica adaptada al sector
        return f"""
        1. ANÁLISIS TÉCNICO ESPECÍFICO
        Evaluación técnica especializada del aspecto a valorar en el contexto del sector {sector},
        identificación de factores críticos y análisis de requerimientos específicos.

        2. PROPUESTA TÉCNICA DIFERENCIADA
        Desarrollo de propuesta técnica adaptada a las especificidades del criterio,
        metodología especializada y enfoque técnico innovador.

        3. IMPLEMENTACIÓN OPERATIVA
        Descripción detallada del proceso de implementación, procedimientos técnicos específicos
        y secuencia operativa optimizada para el sector.

        4. ASEGURAMIENTO DE RESULTADOS
        Sistema de garantía de resultados técnicos, métricas de evaluación específicas
        y protocolos de verificación de cumplimiento.

        5. VALOR AÑADIDO TÉCNICO
        Elementos diferenciadores de la propuesta técnica, innovaciones aplicadas
        y beneficios adicionales aportados al proyecto.
        """

def generar_resumen_personal_tecnico(equipo_tecnico):
    """Genera un resumen estructurado del personal técnico para incluir en prompts de IA"""
    if not equipo_tecnico:
        return "No se ha definido personal técnico específico."

    resumen = ""
    for i, persona in enumerate(equipo_tecnico, 1):
        if persona.get('nombre') or persona.get('cargo'):
            resumen += f"\n{i}. "
            if persona.get('nombre'):
                resumen += f"{persona['nombre']} - "
            resumen += f"{persona.get('cargo', 'Sin especificar')}"
            if persona.get('titulacion'):
                resumen += f" ({persona['titulacion']})"
            resumen += f", {persona.get('experiencia', 0)} años de experiencia"
            if persona.get('certificaciones_personales'):
                resumen += f". Certificaciones: {persona['certificaciones_personales']}"
            if persona.get('experiencia_detalle'):
                experiencia_breve = persona['experiencia_detalle'][:200] + "..." if len(persona['experiencia_detalle']) > 200 else persona['experiencia_detalle']
                resumen += f". Experiencia destacada: {experiencia_breve}"

    return resumen if resumen else "Personal técnico sin detalles especificados."

def generar_seccion_anexos(doc, documentos_anexos):
    """Genera la sección de anexos en el documento Word"""
    try:
        doc.add_heading('ANEXOS', 1)

        # Introducción a los anexos
        intro_p = doc.add_paragraph()
        intro_p.add_run("Los siguientes documentos se adjuntan como anexos para respaldar la propuesta técnica presentada. ")
        intro_p.add_run("Todos los documentos están vigentes y certificados según normativa aplicable.").bold = True

        doc.add_paragraph()  # Espacio

        # Organizar documentos por categoría
        docs_por_categoria = {}
        for doc_info in documentos_anexos:
            categoria = doc_info.get('categoria', 'Otros Documentos')
            if categoria not in docs_por_categoria:
                docs_por_categoria[categoria] = []
            docs_por_categoria[categoria].append(doc_info)

        anexo_num = 1
        for categoria, documentos in docs_por_categoria.items():
            # Título de categoría
            doc.add_heading(f'ANEXO {anexo_num} - {categoria.upper()}', 2)

            # Lista de documentos en esta categoría
            for doc_info in documentos:
                p = doc.add_paragraph()
                p.add_run(f"📄 {doc_info.get('nombre', 'Documento sin nombre')}").bold = True
                p.add_run(f"\nFecha de incorporación: {doc_info.get('fecha_subida', 'N/A')}")

                # Información adicional según el tipo de documento
                if 'ISO' in categoria or 'Certificado' in categoria:
                    p.add_run(f"\nDocumento que acredita el cumplimiento de normativas y estándares de calidad vigentes.")
                elif 'Currículum' in categoria:
                    p.add_run(f"\nDocumentación que respalda la cualificación y experiencia del personal técnico asignado.")
                elif 'Buena Ejecución' in categoria:
                    p.add_run(f"\nCertificado que avala la correcta ejecución de proyectos similares anteriores.")
                else:
                    p.add_run(f"\nDocumento técnico complementario que respalda la propuesta presentada.")

                doc.add_paragraph()  # Espacio entre documentos

            anexo_num += 1

        # Nota final
        doc.add_paragraph()
        nota_p = doc.add_paragraph()
        nota_p.add_run("NOTA IMPORTANTE: ").bold = True
        nota_p.add_run("Todos los documentos anexos están disponibles en formato digital y físico para su verificación. ")
        nota_p.add_run("La documentación se encuentra actualizada a fecha de presentación de esta memoria técnica.")

        doc.add_page_break()

    except Exception as e:
        print(f"Error generando sección de anexos: {e}")
        # Añadir al menos una sección básica si hay error
        doc.add_heading('ANEXOS', 1)
        doc.add_paragraph(f"Se adjuntan {len(documentos_anexos)} documentos como anexos a esta memoria técnica.")
        doc.add_page_break()

def generar_cronograma_proyecto(datos_proyecto, sector='general'):
    """Genera un cronograma profesional según el tipo de proyecto y sector"""
    import plotly.graph_objects as go
    import plotly.express as px
    from datetime import datetime, timedelta
    import pandas as pd

    # Obtener plazo del proyecto
    plazo_str = datos_proyecto.get('plazo', '12 meses')

    # Extraer número de meses del plazo
    try:
        if 'mes' in plazo_str.lower():
            meses = int(''.join(filter(str.isdigit, plazo_str.split('mes')[0])))
        elif 'día' in plazo_str.lower():
            dias = int(''.join(filter(str.isdigit, plazo_str.split('día')[0])))
            meses = max(1, dias // 30)
        else:
            meses = 12  # Default
    except:
        meses = 12

    # Definir fases según sector
    fases_por_sector = {
        'construccion': [
            ('Estudios previos y permisos', 0.15),
            ('Replanteo y preparación', 0.08),
            ('Movimiento de tierras', 0.12),
            ('Cimentación y estructura', 0.25),
            ('Instalaciones', 0.20),
            ('Acabados', 0.15),
            ('Entrega y documentación', 0.05)
        ],
        'electricidad': [
            ('Proyecto ejecutivo y permisos', 0.15),
            ('Suministro de materiales', 0.10),
            ('Montaje de cuadros principales', 0.20),
            ('Cableado e instalación', 0.30),
            ('Conexionado y pruebas', 0.15),
            ('Puesta en marcha', 0.08),
            ('Documentación final', 0.02)
        ],
        'software': [
            ('Análisis de requisitos', 0.15),
            ('Diseño de arquitectura', 0.12),
            ('Desarrollo frontend', 0.25),
            ('Desarrollo backend', 0.25),
            ('Pruebas y testing', 0.15),
            ('Despliegue y configuración', 0.05),
            ('Documentación y entrega', 0.03)
        ],
        'mantenimiento': [
            ('Planificación y programación', 0.10),
            ('Mantenimiento preventivo', 0.40),
            ('Mantenimiento correctivo', 0.25),
            ('Inspecciones y auditorías', 0.15),
            ('Informes y documentación', 0.10)
        ],
        'general': [
            ('Planificación inicial', 0.10),
            ('Fase de preparación', 0.15),
            ('Ejecución principal', 0.50),
            ('Control y seguimiento', 0.15),
            ('Finalización y entrega', 0.10)
        ]
    }

    fases = fases_por_sector.get(sector, fases_por_sector['general'])

    # Crear fechas de inicio y fin para cada fase
    fecha_inicio = datetime.now()
    cronograma_data = []
    fecha_actual = fecha_inicio

    for i, (nombre_fase, porcentaje) in enumerate(fases):
        duracion_dias = int(meses * 30 * porcentaje)
        fecha_fin = fecha_actual + timedelta(days=duracion_dias)

        cronograma_data.append({
            'Fase': nombre_fase,
            'Inicio': fecha_actual,
            'Fin': fecha_fin,
            'Duración': duracion_dias,
            'Porcentaje': porcentaje * 100
        })

        fecha_actual = fecha_fin + timedelta(days=1)

    # Crear DataFrame
    df = pd.DataFrame(cronograma_data)

    # Crear gráfico Gantt
    fig = go.Figure()

    # Colores profesionales para las fases
    colores = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6A994E', '#BC4749', '#F2CC8F']

    for i, row in df.iterrows():
        fig.add_trace(go.Scatter(
            x=[row['Inicio'], row['Fin']],
            y=[i, i],
            mode='lines',
            line=dict(color=colores[i % len(colores)], width=20),
            name=row['Fase'],
            hovertemplate=f"<b>{row['Fase']}</b><br>" +
                         f"Inicio: {row['Inicio'].strftime('%d/%m/%Y')}<br>" +
                         f"Fin: {row['Fin'].strftime('%d/%m/%Y')}<br>" +
                         f"Duración: {row['Duración']} días<br>" +
                         f"Porcentaje: {row['Porcentaje']:.1f}%<extra></extra>"
        ))

    # Personalizar diseño
    fig.update_layout(
        title={
            'text': f"<b>CRONOGRAMA DE EJECUCIÓN - {datos_proyecto.get('objeto', '').upper()}</b>",
            'x': 0.5,
            'xanchor': 'center',
            'font': {'size': 16, 'color': '#2E86AB'}
        },
        xaxis_title="<b>Calendario de Ejecución</b>",
        yaxis_title="<b>Fases del Proyecto</b>",
        yaxis=dict(
            tickmode='array',
            tickvals=list(range(len(df))),
            ticktext=[f"<b>{fase}</b>" for fase in df['Fase']],
            autorange="reversed"
        ),
        xaxis=dict(
            type='date',
            tickformat='%d/%m/%Y'
        ),
        height=400 + len(df) * 30,
        showlegend=False,
        plot_bgcolor='white',
        paper_bgcolor='white',
        font=dict(size=11),
        margin=dict(l=250, r=50, t=80, b=50)
    )

    # Línea vertical para fecha actual comentada para evitar errores de tipo
    # fecha_hoy = datetime.now()
    # fig.add_vline(
    #     x=fecha_hoy,
    #     line_dash="dash",
    #     line_color="red",
    #     annotation_text="HOY",
    #     annotation_position="top"
    # )

    return fig, df

def crear_portada_profesional(doc, datos_proyecto, datos_empresa, logo_path=None):
    """Crea una portada profesional con diseño atractivo"""
    try:
        # Configurar la primera página como portada
        section = doc.sections[0]

        # Título principal centrado
        titulo_principal = doc.add_heading('', 0)
        titulo_run = titulo_principal.runs[0] if titulo_principal.runs else titulo_principal.add_run()
        titulo_run.text = 'MEMORIA TÉCNICA'
        titulo_run.font.size = Pt(24)
        titulo_run.font.bold = True
        titulo_run.font.color.rgb = RGBColor(46, 134, 171)  # Azul profesional
        titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Espacio
        doc.add_paragraph()

        # Crear tabla para estructura visual de la portada
        portada_table = doc.add_table(rows=6, cols=1)
        portada_table.autofit = True

        # Fila 1: Logo (si existe)
        if logo_path and os.path.exists(logo_path):
            logo_cell = portada_table.rows[0].cells[0]
            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            try:
                run.add_picture(logo_path, width=Inches(2.5))
            except:
                run.text = "🏢 LOGO EMPRESA"
                run.font.size = Pt(18)

        # Fila 2: Nombre del proyecto
        proyecto_cell = portada_table.rows[1].cells[0]
        proyecto_para = proyecto_cell.paragraphs[0]
        proyecto_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = proyecto_para.add_run(datos_proyecto.get('objeto', '').upper())
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(33, 37, 41)

        # Fila 3: Información del proyecto
        info_cell = portada_table.rows[2].cells[0]
        info_para = info_cell.paragraphs[0]
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_text = f"Expediente: {datos_proyecto.get('expediente', 'N/A')}\n"
        info_text += f"Organismo: {datos_proyecto.get('organismo', 'N/A')}\n"
        if datos_proyecto.get('presupuesto'):
            info_text += f"Presupuesto: {datos_proyecto.get('presupuesto')} €\n"
        if datos_proyecto.get('plazo'):
            info_text += f"Plazo de ejecución: {datos_proyecto.get('plazo')}"

        run = info_para.add_run(info_text)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(108, 117, 125)

        # Fila 4: Línea separadora visual
        sep_cell = portada_table.rows[3].cells[0]
        sep_para = sep_cell.paragraphs[0]
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep_para.add_run("━" * 50)
        run.font.color.rgb = RGBColor(46, 134, 171)

        # Fila 5: Información de la empresa
        empresa_cell = portada_table.rows[4].cells[0]
        empresa_para = empresa_cell.paragraphs[0]
        empresa_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = empresa_para.add_run("PRESENTADO POR:\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(108, 117, 125)

        run = empresa_para.add_run(f"{datos_empresa.get('razon_social', '')}\n")
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 134, 171)

        if datos_empresa.get('cif'):
            run = empresa_para.add_run(f"CIF: {datos_empresa.get('cif')}\n")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(108, 117, 125)

        # Fila 6: Fecha
        fecha_cell = portada_table.rows[5].cells[0]
        fecha_para = fecha_cell.paragraphs[0]
        fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fecha_para.add_run(f"{datetime.now().strftime('%d de %B de %Y')}")
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(108, 117, 125)

        # Configurar espaciado entre celdas
        for row in portada_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.space_after = Pt(12)
                    paragraph.space_before = Pt(12)

        # Añadir salto de página después de la portada
        doc.add_page_break()

    except Exception as e:
        print(f"Error creando portada: {e}")
        # Portada básica de respaldo
        doc.add_heading('MEMORIA TÉCNICA', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(datos_proyecto.get('objeto', ''), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Expediente: {datos_proyecto.get('expediente', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Empresa: {datos_empresa.get('razon_social', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

def crear_cronograma_tabla_word(doc, df_cronograma, datos_proyecto):
    """Crea una tabla cronograma directamente en Word sin dependencias externas"""
    try:
        # Añadir sección de cronograma al documento
        doc.add_heading('CRONOGRAMA DE EJECUCIÓN', 1)

        doc.add_paragraph(
            f"El siguiente cronograma detalla la planificación temporal propuesta para la ejecución "
            f"del proyecto \"{datos_proyecto.get('objeto', '')}\", con un plazo total de {datos_proyecto.get('plazo', 'N/A')}. "
            f"La planificación ha sido optimizada para garantizar el cumplimiento de los plazos establecidos."
        )

        # Crear tabla del cronograma
        tabla_cronograma = doc.add_table(rows=1, cols=4)
        tabla_cronograma.style = 'Table Grid'

        # Encabezados de tabla
        hdr_cells = tabla_cronograma.rows[0].cells
        hdr_cells[0].text = 'FASE'
        hdr_cells[1].text = 'FECHA INICIO'
        hdr_cells[2].text = 'FECHA FIN'
        hdr_cells[3].text = 'DURACIÓN'

        # Formatear encabezados
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Añadir datos del cronograma
        for _, row in df_cronograma.iterrows():
            row_cells = tabla_cronograma.add_row().cells
            row_cells[0].text = str(row['Fase'])
            row_cells[1].text = row['Inicio'].strftime('%d/%m/%Y')
            row_cells[2].text = row['Fin'].strftime('%d/%m/%Y')
            row_cells[3].text = f"{row['Duración']} días"

            # Formatear celdas
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Añadir gráfico de barras ASCII simple como alternativa visual
        doc.add_paragraph("\nDiagrama temporal:")

        # Calcular duración total para escalar barras
        duracion_total = sum(df_cronograma['Duración'])

        for _, row in df_cronograma.iterrows():
            # Crear barra visual simple con caracteres
            porcentaje = (row['Duración'] / duracion_total) * 100
            barra_largo = int(porcentaje / 2)  # Escalar para que no sea muy larga
            barra = "█" * barra_largo + "░" * (50 - barra_largo)

            p = doc.add_paragraph()
            p.add_run(f"{row['Fase'][:30]:<30} ").font.size = Pt(9)
            p.add_run(f"[{barra[:30]}] ").font.name = 'Courier New'
            p.add_run(f"{porcentaje:.1f}%").font.size = Pt(8)

        return True

    except Exception as e:
        print(f"Error creando cronograma en tabla: {e}")
        return False

def agregar_numeracion_paginas(doc):
    """Añade numeración de páginas al documento"""
    try:
        # Obtener la sección del documento
        section = doc.sections[0]

        # Crear el footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Crear el elemento de numeración
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # Añadir elementos al párrafo del footer
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

        # Aplicar formato
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    except Exception as e:
        print(f"Error añadiendo numeración de páginas: {e}")

def generar_memoria_por_criterios(datos_proyecto, criterios, texto_ppt, datos_empresa):
    """
    Genera contenido profesional desarrollado para cada criterio de valoración
    con estructura adaptada al tipo de criterio y sector específico
    """
    secciones_criterios = {}

    # Obtener análisis avanzado si está disponible
    analisis_ppt = st.session_state.get('analisis_ppt', {})
    sector = detectar_sector_proyecto(datos_proyecto.get('objeto', ''), texto_ppt)

    # Información adicional del análisis avanzado
    requisitos_tecnicos = analisis_ppt.get('requisitos_tecnicos', [])
    tablas_relevantes = [t for t in analisis_ppt.get('tablas', []) if t.get('relevancia', 0) > 60]
    secciones_importantes = analisis_ppt.get('secciones', {})

    for criterio in criterios:
        nombre_criterio = criterio.get('nombre')
        puntos = criterio.get('puntos')

        # Generar estructura específica según el criterio
        estructura_criterio = generar_estructura_especifica(nombre_criterio, sector)

        prompt = f"""
        Actúa como un {get_rol_profesional(sector)} con amplia experiencia en licitaciones públicas.

        Genera una sección técnica especializada para el siguiente criterio de valoración.

        IMPORTANTE:
        - Redacción fluida y profesional en párrafos completos
        - Lenguaje técnico especializado del sector {sector}
        - NO uses listas, viñetas, asteriscos ni símbolos
        - Evita frases de relleno como "Para garantizar el éxito", "La experiencia nos ha enseñado"
        - Incluye especificaciones técnicas detalladas con modelos, marcas y características específicas
        - Profundiza en cada aspecto mencionado con datos técnicos concretos

        CRITERIO: {nombre_criterio}
        PUNTUACIÓN MÁXIMA: {puntos} puntos
        SECTOR: {sector}

        INFORMACIÓN DEL PROYECTO:
        Objeto: {datos_proyecto.get('objeto')}
        Presupuesto: {datos_proyecto.get('presupuesto')} €
        Plazo: {datos_proyecto.get('plazo')}

        INFORMACIÓN DE LA EMPRESA:
        Experiencia: {datos_empresa.get('experiencia')} años
        Certificaciones: {', '.join(datos_empresa.get('certificaciones', []))}

        PERSONAL TÉCNICO ASIGNADO:
        {generar_resumen_personal_tecnico(datos_empresa.get('equipo_tecnico', []))}

        ESTRUCTURA ESPECÍFICA PARA ESTE CRITERIO:
        {estructura_criterio}

        ANÁLISIS AVANZADO DEL PLIEGO (INFORMACIÓN EXTRAÍDA AUTOMÁTICAMENTE):

        REQUISITOS TÉCNICOS IDENTIFICADOS:
        {chr(10).join([f"- {req}" for req in requisitos_tecnicos[:10]])}

        TABLAS RELEVANTES EXTRAÍDAS:
        {chr(10).join([f"- Tabla {t['tipo']}: {len(t['datos'])} elementos (Relevancia: {t['relevancia']}%)" for t in tablas_relevantes[:5]])}

        SECCIONES IMPORTANTES DEL PLIEGO:
        {chr(10).join([f"- {seccion.upper()}: {contenido[:200]}..." if len(contenido) > 200 else f"- {seccion.upper()}: {contenido}" for seccion, contenido in secciones_importantes.items()])}

        CONTEXTO COMPLETO DEL PLIEGO TÉCNICO:
        {texto_ppt[:5000] if texto_ppt else "No disponible"}...

        ESPECIFICACIONES TÉCNICAS DETALLADAS REQUERIDAS:

        Para EQUIPAMIENTO Y RECURSOS TÉCNICOS:
        - Marca, modelo y año de fabricación específicos (Ej: "Caterpillar CAT 320D modelo 2019")
        - Especificaciones técnicas exactas: potencia (kW/CV), capacidad (toneladas), rendimiento (m³/h)
        - Certificaciones CE, normativas ISO aplicables, homologaciones específicas
        - Sistemas de control integrados (GPS, telemetría, sensores)
        - Consumo energético, autonomía operativa, mantenimiento programado

        Para SOFTWARE Y TECNOLOGÍA:
        - Versiones específicas de software (Ej: "AutoCAD 2024, Módulo Civil 3D")
        - Licencias profesionales, compatibilidad, integración con otros sistemas
        - Capacidades de procesamiento, bases de datos, protocolos de seguridad
        - Hardware asociado: servidores (Dell PowerEdge R740xd, 64GB RAM, Intel Xeon)
        - Sistemas de backup, redundancia, conectividad 5G/fibra óptica

        Para INFRAESTRUCTURA Y INSTALACIONES:
        - Ubicación física exacta, superficie (m²), distribución funcional
        - Sistemas de climatización (Mitsubishi Electric VRF, 25kW)
        - Conectividad: fibra óptica simétrica 1Gbps, redes redundantes
        - Sistemas de seguridad: videovigilancia 4K, control accesos biométrico
        - Certificaciones energéticas, sistemas UPS (APC Smart-UPS RT 10kVA)

        Para METODOLOGÍAS Y PROCESOS:
        - Protocolos específicos según normativa (UNE, ISO, AENOR)
        - Procedimientos certificados, check-lists técnicos detallados
        - Sistemas de trazabilidad: códigos QR, RFID, blockchain para verificación
        - Indicadores KPI específicos con umbrales de aceptación numéricos
        - Herramientas de medición calibradas (certificado de calibración vigente)

        INTEGRACIÓN DE PERSONAL TÉCNICO EN CONTENIDO:
        - Menciona nombres específicos del equipo cuando sea relevante técnicamente
        - Relaciona titulaciones con competencias específicas del criterio
        - Referencias a experiencia previa detallada en proyectos similares
        - Certificaciones personales aplicables al criterio específico

        REQUISITOS DE CALIDAD TÉCNICA:
        - Mínimo 3500 palabras por criterio con contenido técnico sustancial
        - Párrafos de 180-280 palabras con alta densidad técnica
        - Eliminación total de frases de relleno y generalidades
        - Datos técnicos verificables con referencias normativas específicas
        - Lenguaje especializado del sector con precisión terminológica
        """
        
        # Para mejor rendimiento, usar GPT-4 o Claude
        respuesta = llamar_ia_mejorado(prompt, max_tokens=4000, temperature=0.3)
        
        # Limpieza adicional para eliminar cualquier símbolo no deseado
        if respuesta:
            respuesta = respuesta.replace('*', '')
            respuesta = respuesta.replace('#', '')
            respuesta = respuesta.replace('•', '')
            respuesta = respuesta.replace('→', '')
            # Eliminar listas numeradas
            respuesta = re.sub(r'^\d+\.?\s+', '', respuesta, flags=re.MULTILINE)
            # Eliminar guiones al inicio de línea
            respuesta = re.sub(r'^-\s+', '', respuesta, flags=re.MULTILINE)
            
        secciones_criterios[nombre_criterio] = respuesta
    
    return secciones_criterios

def extraer_texto_pdf_avanzado(archivo_pdf):
    """Extrae texto, tablas y estructura de PDF con análisis avanzado"""
    try:
        import pdfplumber
        texto_completo = ""
        tablas_extraidas = []
        criterios_detectados = []
        requisitos_tecnicos = []
        secciones_importantes = {}

        with pdfplumber.open(archivo_pdf) as pdf:
            for num_pagina, page in enumerate(pdf.pages, 1):
                # Extraer texto de la página
                texto_pagina = page.extract_text()
                if texto_pagina:
                    texto_completo += f"\n--- PÁGINA {num_pagina} ---\n"
                    texto_completo += texto_pagina + "\n"

                    # Analizar contenido de la página
                    analizar_contenido_pagina(texto_pagina, criterios_detectados, requisitos_tecnicos, secciones_importantes)

                # Extraer tablas de la página
                tablas = page.extract_tables()
                if tablas:
                    for i, tabla in enumerate(tablas):
                        if tabla and len(tabla) > 1:  # Verificar que la tabla tenga contenido
                            tabla_procesada = procesar_tabla_extraida(tabla, num_pagina, i)
                            if tabla_procesada:
                                tablas_extraidas.append(tabla_procesada)
                                texto_completo += f"\n--- TABLA {num_pagina}.{i+1} ---\n"
                                texto_completo += formatear_tabla_como_texto(tabla_procesada)

        # Crear análisis estructurado
        analisis_completo = {
            'texto_completo': texto_completo,
            'tablas': tablas_extraidas,
            'criterios_valoracion': criterios_detectados,
            'requisitos_tecnicos': requisitos_tecnicos,
            'secciones': secciones_importantes,
            'resumen_analisis': generar_resumen_analisis(criterios_detectados, requisitos_tecnicos)
        }

        return analisis_completo

    except ImportError:
        st.warning("⚠️ pdfplumber no disponible. Usando extracción básica...")
        return extraer_texto_pdf_basico(archivo_pdf)
    except Exception as e:
        st.error(f"Error procesando PDF: {str(e)}")
        return extraer_texto_pdf_basico(archivo_pdf)

def extraer_texto_pdf_basico(archivo_pdf):
    """Versión básica de extracción (fallback)"""
    try:
        pdf_reader = PyPDF2.PdfReader(archivo_pdf)
        texto = ""
        for page in pdf_reader.pages:
            texto += page.extract_text() + "\n"
        return {'texto_completo': texto, 'tablas': [], 'criterios_valoracion': [], 'requisitos_tecnicos': [], 'secciones': {}}
    except Exception as e:
        st.error(f"Error procesando PDF: {str(e)}")
        return {'texto_completo': "", 'tablas': [], 'criterios_valoracion': [], 'requisitos_tecnicos': [], 'secciones': {}}

def analizar_contenido_pagina(texto_pagina, criterios_detectados, requisitos_tecnicos, secciones_importantes):
    """Analiza el contenido de una página para extraer información estructurada"""
    texto_lower = texto_pagina.lower()

    # Detectar criterios de valoración
    patrones_criterios = [
        r'criterio\s+(?:de\s+)?valoraci[óo]n[:\s]*(.{0,100})',
        r'criterio[:\s]*(.{0,100}?)(?:\d+\s*puntos?|\d+\s*%)',
        r'puntuaci[óo]n[:\s]*(.{0,100}?)(?:\d+\s*puntos?)',
        r'valoraci[óo]n[:\s]*(.{0,100}?)(?:\d+\s*puntos?)',
        r'apartado[:\s]*(.{0,100}?)(?:\d+\s*puntos?)'
    ]

    for patron in patrones_criterios:
        matches = re.findall(patron, texto_lower, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            criterio_limpio = limpiar_criterio_detectado(match)
            if criterio_limpio and len(criterio_limpio) > 10:
                puntos = extraer_puntos_criterio(texto_pagina, match)
                criterios_detectados.append({
                    'nombre': criterio_limpio,
                    'puntos': puntos,
                    'texto_original': match
                })

    # Detectar requisitos técnicos específicos
    patrones_requisitos = [
        r'requisito[s]?\s+t[ée]cnico[s]?[:\s]*(.{0,200})',
        r'especificaci[óo]n[es]?\s+t[ée]cnica[s]?[:\s]*(.{0,200})',
        r'caracter[íi]stica[s]?\s+t[ée]cnica[s]?[:\s]*(.{0,200})',
        r'normativa[s]?\s+aplicable[s]?[:\s]*(.{0,200})',
        r'certificaci[óo]n[es]?\s+requerida[s]?[:\s]*(.{0,200})'
    ]

    for patron in patrones_requisitos:
        matches = re.findall(patron, texto_lower, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            requisito_limpio = match.strip()
            if len(requisito_limpio) > 15:
                requisitos_tecnicos.append(requisito_limpio)

    # Detectar secciones importantes por títulos
    secciones_titulos = [
        'objeto del contrato', 'criterios de valoración', 'requisitos técnicos',
        'especificaciones técnicas', 'condiciones de ejecución', 'garantías',
        'personal técnico', 'medios materiales', 'experiencia'
    ]

    for seccion in secciones_titulos:
        if seccion in texto_lower:
            # Extraer contenido después del título
            patron_seccion = rf'{re.escape(seccion)}[:\s]*(.{{0,500}}?)(?:\n\n|\n[A-Z])'
            match = re.search(patron_seccion, texto_pagina, re.IGNORECASE | re.DOTALL)
            if match:
                secciones_importantes[seccion] = match.group(1).strip()

def procesar_tabla_extraida(tabla, num_pagina, indice_tabla):
    """Procesa una tabla extraída identificando su tipo y contenido relevante"""
    if not tabla or len(tabla) < 2:
        return None

    # Limpiar celdas vacías
    tabla_limpia = []
    for fila in tabla:
        fila_limpia = [str(celda).strip() if celda else "" for celda in fila]
        if any(celda for celda in fila_limpia):  # Solo agregar filas no vacías
            tabla_limpia.append(fila_limpia)

    if len(tabla_limpia) < 2:
        return None

    # Identificar tipo de tabla
    headers = tabla_limpia[0]
    tipo_tabla = identificar_tipo_tabla(headers)

    return {
        'pagina': num_pagina,
        'indice': indice_tabla,
        'tipo': tipo_tabla,
        'headers': headers,
        'datos': tabla_limpia[1:],
        'relevancia': calcular_relevancia_tabla(headers, tabla_limpia)
    }

def identificar_tipo_tabla(headers):
    """Identifica el tipo de tabla basado en sus encabezados"""
    headers_lower = [str(h).lower() for h in headers if h]

    if any(word in ' '.join(headers_lower) for word in ['criterio', 'puntos', 'valoración']):
        return 'criterios_valoracion'
    elif any(word in ' '.join(headers_lower) for word in ['requisito', 'especificación', 'técnico']):
        return 'requisitos_tecnicos'
    elif any(word in ' '.join(headers_lower) for word in ['plazo', 'fecha', 'entrega']):
        return 'plazos'
    elif any(word in ' '.join(headers_lower) for word in ['presupuesto', 'importe', 'precio']):
        return 'presupuesto'
    else:
        return 'general'

def calcular_relevancia_tabla(headers, datos):
    """Calcula la relevancia de la tabla para la memoria técnica"""
    palabras_relevantes = ['criterio', 'puntos', 'valoración', 'técnico', 'requisito', 'especificación']
    contenido_completo = ' '.join(str(item) for fila in [headers] + datos for item in fila if item).lower()

    relevancia = sum(contenido_completo.count(palabra) for palabra in palabras_relevantes)
    return min(relevancia * 10, 100)  # Normalizar a 0-100

def formatear_tabla_como_texto(tabla_procesada):
    """Convierte una tabla en formato texto legible"""
    texto_tabla = f"\nTipo: {tabla_procesada['tipo'].replace('_', ' ').title()}\n"
    texto_tabla += f"Relevancia: {tabla_procesada['relevancia']}%\n\n"

    headers = tabla_procesada['headers']
    datos = tabla_procesada['datos']

    # Crear formato tabular
    for i, header in enumerate(headers):
        texto_tabla += f"{header:<20}"
    texto_tabla += "\n" + "-" * (len(headers) * 20) + "\n"

    for fila in datos[:10]:  # Limitar a 10 filas para evitar exceso
        for i, celda in enumerate(fila):
            if i < len(headers):
                texto_tabla += f"{str(celda)[:18]:<20}"
        texto_tabla += "\n"

    return texto_tabla

def limpiar_criterio_detectado(criterio_raw):
    """Limpia y normaliza un criterio detectado"""
    criterio = criterio_raw.strip()
    # Eliminar caracteres especiales al inicio/final
    criterio = re.sub(r'^[:\-\.\s]+|[:\-\.\s]+$', '', criterio)
    # Capitalizar primera letra
    if criterio:
        criterio = criterio[0].upper() + criterio[1:]
    return criterio

def extraer_puntos_criterio(texto, criterio):
    """Extrae los puntos asignados a un criterio"""
    # Buscar patrones de puntuación cerca del criterio
    patron_puntos = r'(\d+)\s*(?:puntos?|pts?|%)'

    # Buscar en un contexto de 200 caracteres alrededor del criterio
    inicio = max(0, texto.lower().find(criterio.lower()) - 100)
    fin = min(len(texto), texto.lower().find(criterio.lower()) + len(criterio) + 100)
    contexto = texto[inicio:fin]

    matches = re.findall(patron_puntos, contexto, re.IGNORECASE)
    return int(matches[0]) if matches else 0

def generar_resumen_analisis(criterios, requisitos):
    """Genera un resumen del análisis del pliego"""
    resumen = f"""
ANÁLISIS AUTOMÁTICO DEL PLIEGO TÉCNICO:

📊 CRITERIOS DE VALORACIÓN DETECTADOS: {len(criterios)}
{chr(10).join([f"• {c['nombre']} ({c['puntos']} puntos)" for c in criterios[:10]])}

🔧 REQUISITOS TÉCNICOS IDENTIFICADOS: {len(requisitos)}
{chr(10).join([f"• {req[:100]}..." if len(req) > 100 else f"• {req}" for req in requisitos[:10]])}

📈 NIVEL DE ANÁLISIS: AVANZADO
✅ Extracción de tablas: Activa
✅ Detección de criterios: Automática
✅ Análisis estructural: Completo
"""
    return resumen

def extraer_datos_licitacion(url_licitacion):
    """
    Extrae datos automáticamente de la Plataforma de Contratación del Sector Público
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Hacer la petición
        response = requests.get(url_licitacion, headers=headers, timeout=15, verify=False)
        response.encoding = 'utf-8'
        soup = BeautifulSoup(response.text, 'html.parser')
        
        datos_extraidos = {}
        
        # Estrategia 1: Buscar en tablas con class "table-view"
        tablas = soup.find_all('table', class_='table-view')
        for tabla in tablas:
            filas = tabla.find_all('tr')
            for fila in filas:
                celdas = fila.find_all('td')
                if len(celdas) >= 2:
                    etiqueta = celdas[0].get_text(strip=True)
                    valor = celdas[1].get_text(strip=True)
                    
                    if 'Objeto' in etiqueta and 'objeto' not in datos_extraidos:
                        datos_extraidos['objeto'] = valor
                    elif 'Presupuesto base' in etiqueta and 'sin IVA' in etiqueta:
                        # Extraer solo números
                        numeros = re.findall(r'[\d.,]+', valor)
                        if numeros:
                            datos_extraidos['presupuesto'] = numeros[0]
                    elif 'Número de expediente' in etiqueta:
                        datos_extraidos['expediente'] = valor
                    elif 'Plazo' in etiqueta or 'Duración' in etiqueta:
                        datos_extraidos['plazo'] = valor
                    elif 'Órgano de Contratación' in etiqueta:
                        datos_extraidos['organismo'] = valor
        
        # Estrategia 2: Buscar en divs con class "detail-info"
        if not datos_extraidos.get('objeto'):
            detalles = soup.find_all('div', class_='detail-info')
            for detalle in detalles:
                titulo = detalle.find('span', class_='detail-title')
                valor = detalle.find('span', class_='detail-value')
                if titulo and valor:
                    titulo_text = titulo.get_text(strip=True)
                    valor_text = valor.get_text(strip=True)
                    
                    if 'Objeto' in titulo_text:
                        datos_extraidos['objeto'] = valor_text
                    elif 'Expediente' in titulo_text:
                        datos_extraidos['expediente'] = valor_text
        
        # Estrategia 3: Buscar en el contenido principal
        if not datos_extraidos.get('objeto'):
            # Buscar por ID o class específicos
            objeto_elem = soup.find(id=re.compile('.*objeto.*', re.I))
            if objeto_elem:
                datos_extraidos['objeto'] = objeto_elem.get_text(strip=True)
            
            # Buscar presupuesto
            presupuesto_elem = soup.find(text=re.compile(r'[\d.,]+ EUR', re.I))
            if presupuesto_elem:
                match = re.search(r'([\d.,]+)\s*EUR', presupuesto_elem)
                if match:
                    datos_extraidos['presupuesto'] = match.group(1)
        
        # Estrategia 4: Buscar en toda la página con expresiones regulares
        if not datos_extraidos.get('objeto'):
            texto_completo = soup.get_text()
            
            # Buscar objeto del contrato
            objeto_match = re.search(r'Objeto del contrato[:\s]+([^\n]{10,200})', texto_completo, re.I)
            if objeto_match:
                datos_extraidos['objeto'] = objeto_match.group(1).strip()
            
            # Buscar número de expediente
            exp_match = re.search(r'(?:Expediente|Referencia)[:\s]+([A-Z0-9\-/]+)', texto_completo, re.I)
            if exp_match:
                datos_extraidos['expediente'] = exp_match.group(1).strip()
            
            # Buscar presupuesto
            pres_match = re.search(r'Presupuesto.*?sin IVA[:\s]+([\d.,]+)', texto_completo, re.I)
            if pres_match:
                datos_extraidos['presupuesto'] = pres_match.group(1).strip()
        
        # Si no encontramos datos, intentar con JavaScript renderizado
        if not datos_extraidos:
            st.warning("⚠️ No se pudieron extraer datos automáticamente. Intenta copiar y pegar los datos manualmente.")
            return {}
        
        # Mensaje de éxito con los datos encontrados
        encontrados = [k for k, v in datos_extraidos.items() if v]
        if encontrados:
            st.success(f"✅ Datos extraídos: {', '.join(encontrados)}")
        
        return datos_extraidos
        
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Error de conexión: {str(e)}")
        st.info("💡 Intenta copiar la URL directamente desde el navegador")
        return {}
    except Exception as e:
        st.error(f"❌ Error al procesar la página: {str(e)}")
        st.info("💡 Puedes introducir los datos manualmente")
        return {}

def buscar_info_empresa(razon_social, cif=None):
    """Busca información de la empresa en internet"""
    try:
        prompt = f"""
        Basándote en tu conocimiento sobre la empresa "{razon_social}" 
        {f'con CIF {cif}' if cif else ''}, genera una presentación profesional.
        
        Incluye:
        1. Historia y trayectoria
        2. Principales líneas de negocio
        3. Certificaciones del sector
        4. Datos económicos aproximados
        5. Proyectos destacados
        
        Formato: JSON estructurado
        """
        
        respuesta = llamar_ia_mejorado(prompt, max_tokens=1500, temperature=0.7)
        
        try:
            return json.loads(respuesta)
        except:
            return {"descripcion": respuesta}
    except:
        return None

def calcular_extension_contenido(num_paginas):
    """Calcula la extensión del contenido basado en el número de páginas"""
    extensiones = {
        10: {"palabras": 3000, "tokens": 1500, "detalle": "muy resumido"},
        20: {"palabras": 6000, "tokens": 2500, "detalle": "resumido"},
        30: {"palabras": 9000, "tokens": 3500, "detalle": "normal"},
        40: {"palabras": 12000, "tokens": 4500, "detalle": "detallado"},
        50: {"palabras": 15000, "tokens": 5500, "detalle": "muy detallado"},
        80: {"palabras": 24000, "tokens": 8000, "detalle": "exhaustivo"},
        120: {"palabras": 36000, "tokens": 12000, "detalle": "máximo detalle"}
    }
    
    for paginas, config in extensiones.items():
        if num_paginas <= paginas:
            return config
    
    return extensiones[120]

# ============ PANEL DE ADMINISTRACIÓN ============
def mostrar_aplicacion_admin():
    """Panel de administración para gestionar usuarios y pagos"""
    aplicar_estilos_app()

    # Header administrativo con logo centrado
    col_logo1, col_logo2, col_logo3 = st.columns([2, 1, 2])
    with col_logo2:
        try:
            st.image("logo.png", width=150)
        except:
            pass

    st.markdown(f"""
    <div class="main-header">
        <span class="logo-memoria">{get_text('admin_panel')}</span>
        <h2 style="margin: 0; font-weight: 300;">{get_text('user_management')} - MEMOR.IA</h2>
    </div>
    """, unsafe_allow_html=True)

    # Info de admin
    col1, col2 = st.columns([6, 1])
    with col2:
        st.markdown(f"""
        <div style="text-align: right; padding: 10px;">
            <small>👤 {st.session_state.user_data['nombre']}<br>
            🔑 Administrador</small>
        </div>
        """, unsafe_allow_html=True)

    # Tabs del panel admin
    tab1, tab2, tab3, tab4 = st.tabs([
        "👥 Gestión de Usuarios",
        "💳 Control de Pagos",
        "📊 Estadísticas",
        "🧠 Generar Memoria"
    ])

    with tab1:
        st.markdown('<h2 class="section-header">Gestión de Usuarios Registrados</h2>', unsafe_allow_html=True)

        # Obtener todos los usuarios
        conn = sqlite3.connect('memoria_usuarios.db')
        df_usuarios = pd.read_sql_query('''
            SELECT id, email, nombre, empresa, telefono, cif, numero_cuenta, rol,
                   fecha_registro, activo, plan, fecha_expiracion
            FROM usuarios ORDER BY fecha_registro DESC
        ''', conn)
        conn.close()

        # Filtros
        col_f1, col_f2, col_f3 = st.columns(3)

        with col_f1:
            filtro_activo = st.selectbox("Estado", ["Todos", "Activos", "Inactivos"])

        with col_f2:
            filtro_plan = st.selectbox("Plan", ["Todos", "basico", "premium"])

        with col_f3:
            buscar_email = st.text_input("Buscar por email", placeholder="usuario@email.com")

        # Aplicar filtros
        df_filtrado = df_usuarios.copy()

        if filtro_activo == "Activos":
            df_filtrado = df_filtrado[df_filtrado['activo'] == 1]
        elif filtro_activo == "Inactivos":
            df_filtrado = df_filtrado[df_filtrado['activo'] == 0]

        if filtro_plan != "Todos":
            df_filtrado = df_filtrado[df_filtrado['plan'] == filtro_plan]

        if buscar_email:
            df_filtrado = df_filtrado[df_filtrado['email'].str.contains(buscar_email, case=False)]

        # Mostrar estadísticas rápidas
        col_s1, col_s2, col_s3, col_s4 = st.columns(4)

        with col_s1:
            st.metric("Total Usuarios", len(df_usuarios))
        with col_s2:
            st.metric("Usuarios Activos", len(df_usuarios[df_usuarios['activo'] == 1]))
        with col_s3:
            st.metric("Usuarios Inactivos", len(df_usuarios[df_usuarios['activo'] == 0]))
        with col_s4:
            st.metric("Nuevos Hoy", len(df_usuarios[df_usuarios['fecha_registro'].str.contains(datetime.now().strftime('%Y-%m-%d'))]))

        # Tabla de usuarios
        st.markdown("### 📋 Lista de Usuarios")

        if len(df_filtrado) > 0:
            # Configurar columnas para mostrar
            columnas_mostrar = ['email', 'nombre', 'empresa', 'plan', 'fecha_registro', 'activo']

            for idx, row in df_filtrado.iterrows():
                with st.expander(f"👤 {row['nombre']} - {row['email']} ({'✅ Activo' if row['activo'] else '❌ Inactivo'})"):
                    col_u1, col_u2 = st.columns(2)

                    with col_u1:
                        st.write(f"**Empresa:** {row['empresa']}")
                        st.write(f"**Teléfono:** {row['telefono']}")
                        st.write(f"**CIF:** {row['cif']}")
                        st.write(f"**Plan:** {row['plan']}")

                    with col_u2:
                        st.write(f"**Registro:** {row['fecha_registro']}")
                        st.write(f"**Expiración:** {row['fecha_expiracion']}")
                        st.write(f"**Estado:** {'Activo' if row['activo'] else 'Inactivo'}")

                    # Mostrar número de cuenta en una fila separada
                    st.markdown("---")
                    st.write(f"💳 **Número de Cuenta:** {row['numero_cuenta'] if row['numero_cuenta'] else 'No proporcionado'}")

                    # Acciones
                    col_a1, col_a2, col_a3 = st.columns(3)

                    with col_a1:
                        nuevo_estado = not row['activo']
                        if st.button(f"{'Desactivar' if row['activo'] else 'Activar'}",
                                   key=f"toggle_{row['id']}"):
                            conn = sqlite3.connect('memoria_usuarios.db')
                            cursor = conn.cursor()
                            cursor.execute('''
                                UPDATE usuarios SET activo = ? WHERE id = ?
                            ''', (nuevo_estado, row['id']))
                            conn.commit()
                            conn.close()
                            st.success(f"Usuario {'activado' if nuevo_estado else 'desactivado'}")
                            st.rerun()

                    with col_a2:
                        if st.button("Enviar credenciales", key=f"send_{row['id']}"):
                            # Generar nueva contraseña
                            nueva_pass = generar_password()
                            password_hash = hashlib.sha256(nueva_pass.encode()).hexdigest()

                            conn = sqlite3.connect('memoria_usuarios.db')
                            cursor = conn.cursor()
                            cursor.execute('''
                                UPDATE usuarios SET password = ? WHERE id = ?
                            ''', (password_hash, row['id']))
                            conn.commit()
                            conn.close()

                            # Enviar email
                            mensaje_html = f"""
                            <html>
                            <body>
                                <h2>Nuevas credenciales - MEMOR.IA</h2>
                                <p>Hola {row['nombre']},</p>

                                <h3>Tus nuevas credenciales son:</h3>
                                <p><strong>Email:</strong> {row['email']}</p>
                                <p><strong>Contraseña:</strong> {nueva_pass}</p>

                                <p>Accede en: <a href="http://localhost:8502">MEMOR.IA</a></p>

                                <p><em>Equipo MEMOR.IA</em></p>
                            </body>
                            </html>
                            """

                            if enviar_email(row['email'], "Nuevas credenciales - MEMOR.IA", mensaje_html):
                                st.success("Credenciales enviadas")
                            else:
                                st.error("Error enviando credenciales")

                    with col_a3:
                        if st.button("🗑️ Eliminar", key=f"delete_{row['id']}"):
                            if st.session_state.get(f'confirm_delete_{row["id"]}', False):
                                conn = sqlite3.connect('memoria_usuarios.db')
                                cursor = conn.cursor()
                                cursor.execute('DELETE FROM usuarios WHERE id = ?', (row['id'],))
                                cursor.execute('DELETE FROM pagos WHERE usuario_id = ?', (row['id'],))
                                conn.commit()
                                conn.close()
                                st.success("Usuario eliminado")
                                st.rerun()
                            else:
                                st.session_state[f'confirm_delete_{row["id"]}'] = True
                                st.warning("⚠️ Confirma eliminación haciendo clic nuevamente")
        else:
            st.info("No se encontraron usuarios con los filtros aplicados")

    with tab2:
        st.markdown('<h2 class="section-header">Control de Pagos y Facturación</h2>', unsafe_allow_html=True)

        # Obtener pagos
        conn = sqlite3.connect('memoria_usuarios.db')
        df_pagos = pd.read_sql_query('''
            SELECT p.*, u.nombre, u.email, u.empresa
            FROM pagos p
            JOIN usuarios u ON p.usuario_id = u.id
            ORDER BY p.fecha_pago DESC
        ''', conn)
        conn.close()

        if len(df_pagos) > 0:
            # Métricas de pagos
            col_p1, col_p2, col_p3, col_p4 = st.columns(4)

            total_ingresos = df_pagos['importe'].sum()
            pagos_mes = df_pagos[df_pagos['fecha_pago'].str.contains(datetime.now().strftime('%Y-%m'))]

            with col_p1:
                st.metric("Total Ingresos", f"{total_ingresos:.2f}€")
            with col_p2:
                st.metric("Pagos Este Mes", len(pagos_mes))
            with col_p3:
                st.metric("Ingresos Este Mes", f"{pagos_mes['importe'].sum():.2f}€")
            with col_p4:
                st.metric("Pago Promedio", f"{df_pagos['importe'].mean():.2f}€")

            # Lista de pagos
            st.markdown("### 💳 Historial de Pagos")

            for idx, row in df_pagos.iterrows():
                with st.expander(f"💰 {row['nombre']} - {row['importe']}€ - {row['fecha_pago'][:10]}"):
                    col_p1, col_p2 = st.columns(2)

                    with col_p1:
                        st.write(f"**Cliente:** {row['nombre']}")
                        st.write(f"**Email:** {row['email']}")
                        st.write(f"**Empresa:** {row['empresa']}")
                        st.write(f"**Plan:** {row['plan']}")

                    with col_p2:
                        st.write(f"**Importe:** {row['importe']}€")
                        st.write(f"**Fecha:** {row['fecha_pago']}")
                        st.write(f"**Estado:** {row['estado']}")
                        st.write(f"**Stripe ID:** {row['stripe_payment_id']}")

                    # Opciones de gestión de pagos
                    if row['estado'] == 'pendiente':
                        if st.button("Marcar como Pagado", key=f"mark_paid_{row['id']}"):
                            conn = sqlite3.connect('memoria_usuarios.db')
                            cursor = conn.cursor()
                            cursor.execute('''
                                UPDATE pagos SET estado = 'completado' WHERE id = ?
                            ''', (row['id'],))
                            conn.commit()
                            conn.close()
                            st.success("Pago marcado como completado")
                            st.rerun()
        else:
            st.info("No hay registros de pagos aún")

    with tab3:
        st.markdown('<h2 class="section-header">Estadísticas del Sistema</h2>', unsafe_allow_html=True)

        # Gráfico de registros por fecha
        if len(df_usuarios) > 0:
            df_usuarios['fecha'] = pd.to_datetime(df_usuarios['fecha_registro']).dt.date
            registros_por_dia = df_usuarios.groupby('fecha').size().reset_index(name='registros')

            fig = go.Figure(data=go.Scatter(
                x=registros_por_dia['fecha'],
                y=registros_por_dia['registros'],
                mode='lines+markers',
                name='Registros por día'
            ))
            fig.update_layout(
                title="Evolución de Registros",
                xaxis_title="Fecha",
                yaxis_title="Nuevos Registros"
            )
            st.plotly_chart(fig, use_container_width=True)

            # Distribución por planes
            col_g1, col_g2 = st.columns(2)

            with col_g1:
                plan_counts = df_usuarios['plan'].value_counts()
                fig_pie = go.Figure(data=[go.Pie(
                    labels=plan_counts.index,
                    values=plan_counts.values
                )])
                fig_pie.update_layout(title="Distribución por Planes")
                st.plotly_chart(fig_pie, use_container_width=True)

            with col_g2:
                estado_counts = df_usuarios['activo'].map({1: 'Activos', 0: 'Inactivos'}).value_counts()
                fig_bar = go.Figure(data=[go.Bar(
                    x=estado_counts.index,
                    y=estado_counts.values
                )])
                fig_bar.update_layout(title="Estado de Usuarios")
                st.plotly_chart(fig_bar, use_container_width=True)

    with tab4:
        st.info("🧠 **Panel de Administrador:** También puedes generar memorias técnicas desde aquí")
        mostrar_aplicacion()

# ============ APLICACIÓN PRINCIPAL ============
def mostrar_aplicacion():
    aplicar_estilos_app()
    
    # Header con logo centrado
    col_logo1, col_logo2, col_logo3 = st.columns([2, 1, 2])
    with col_logo2:
        try:
            st.image("logo.png", width=150)
        except:
            pass

    st.markdown(f"""
    <div class="main-header">
        <span class="logo-memoria">MEMOR.IA</span>
        <h2 style="margin: 0; font-weight: 300;">{get_text('memory_generator')}</h2>
    </div>
    """, unsafe_allow_html=True)

    # Advertencia modo demo
    if st.session_state.get('demo_mode', False):
        st.warning("🎭 **MODO DEMOSTRACIÓN** - Esta es una vista previa. Para usar todas las funciones, regístrate como cliente.")

    # Info de usuario
    col1, col2 = st.columns([6, 1])
    with col2:
        st.markdown(f"""
        <div style="text-align: right; padding: 10px;">
            <small>👤 {st.session_state.user_data['nombre']}<br>
            🏢 {st.session_state.user_data['empresa']}</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Inicializar variables en session_state
    if 'datos_extraidos' not in st.session_state:
        st.session_state.datos_extraidos = {}
    if 'texto_ppt' not in st.session_state:
        st.session_state.texto_ppt = ""
    if 'info_empresa_online' not in st.session_state:
        st.session_state.info_empresa_online = None
    if 'criterios_valoracion' not in st.session_state:
        st.session_state.criterios_valoracion = []
    
    # Tabs principales
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "🏢 Perfil de Empresa",
        "📄 Extracción Automática",
        "📝 Datos del Proyecto",
        "🎯 Criterios de Valoración",
        "🔧 Configuración Técnica",
        "📄 Generar Memoria"
    ])
    
    with tab1:
        st.markdown('<h2 class="section-header">🏢 Perfil de Empresa</h2>', unsafe_allow_html=True)

        st.info("💡 **Importante:** Completa tu perfil de empresa una vez. Los datos se cargarán automáticamente en todas tus memorias técnicas.")

        # Obtener perfil existente
        perfil_actual = obtener_perfil_empresa(st.session_state.user_email)

        with st.form("perfil_empresa_form"):
            st.markdown("### 📊 Información General")

            col1, col2 = st.columns(2)

            with col1:
                sector = st.selectbox("Sector", [
                    "Construcción", "Instalaciones Eléctricas", "Energía Fotovoltaica",
                    "Climatización y Ventilación", "Fontanería y Saneamiento",
                    "Ingeniería Civil", "Consultoría Técnica", "Servicios Industriales",
                    "Formación", "Desarrollo de Software", "Ingeniería y Arquitectura",
                    "Limpieza y Jardinería", "Suministros"
                ], index=0 if not perfil_actual else ([
                    "Construcción", "Instalaciones Eléctricas", "Energía Fotovoltaica",
                    "Climatización y Ventilación", "Fontanería y Saneamiento",
                    "Ingeniería Civil", "Consultoría Técnica", "Servicios Industriales",
                    "Formación", "Desarrollo de Software", "Ingeniería y Arquitectura",
                    "Limpieza y Jardinería", "Suministros"
                ].index(perfil_actual['sector']) if perfil_actual['sector'] in [
                    "Construcción", "Instalaciones Eléctricas", "Energía Fotovoltaica",
                    "Climatización y Ventilación", "Fontanería y Saneamiento",
                    "Ingeniería Civil", "Consultoría Técnica", "Servicios Industriales",
                    "Formación", "Desarrollo de Software", "Ingeniería y Arquitectura",
                    "Limpieza y Jardinería", "Suministros"
                ] else 0))

                empleados = st.text_input("Número de empleados",
                                        value=perfil_actual['empleados'] if perfil_actual else "20-50")

                experiencia = st.text_input("Años de experiencia",
                                          value=perfil_actual['experiencia_anos'] if perfil_actual else "10")

            with col2:
                certificaciones = st.multiselect("Certificaciones",
                    ["ISO 9001:2015", "ISO 14001:2015", "ISO 45001:2018",
                     "OHSAS 18001", "ISO 50001", "Marca CE", "ISO 27001",
                     "ENS", "Plan de Igualdad", "Huella de Carbono"],
                    default=perfil_actual['certificaciones'] if perfil_actual else [])

                otras_certificaciones = st.text_input("Otras certificaciones",
                    value=perfil_actual['otras_certificaciones'] if perfil_actual else "",
                    placeholder="Especifica otras certificaciones...")

            # Logo de empresa
            st.markdown("### 🎨 Logo de Empresa")
            col_logo1, col_logo2 = st.columns([1, 2])

            with col_logo1:
                logo_file = st.file_uploader(
                    "Logo (PNG/JPG)",
                    type=['png', 'jpg', 'jpeg'],
                    help="Se usará en todas las memorias técnicas"
                )

            with col_logo2:
                if perfil_actual and perfil_actual.get('logo_path'):
                    try:
                        if os.path.exists(perfil_actual['logo_path']):
                            st.image(perfil_actual['logo_path'], width=150, caption="Logo actual")
                            st.success("✅ Logo guardado correctamente")
                        else:
                            st.warning(f"⚠️ Logo guardado pero archivo no encontrado en: {perfil_actual['logo_path']}")
                    except Exception as e:
                        st.error(f"Error mostrando logo: {e}")
                        st.info("Logo guardado (error al mostrar)")
                else:
                    st.info("No hay logo guardado")

            # Experiencia
            st.markdown("### 📊 Experiencia en Proyectos")
            experiencia_similar = st.text_area(
                "Experiencia en proyectos similares",
                value=perfil_actual['experiencia_similar'] if perfil_actual else "",
                placeholder="Describe la experiencia en proyectos similares...",
                height=100
            )

            # Medios técnicos
            st.markdown("### 🔧 Medios Técnicos")
            col_medios1, col_medios2 = st.columns(2)

            with col_medios1:
                medios_materiales = st.text_area(
                    "Maquinaria y Equipos",
                    value=perfil_actual['medios_materiales'] if perfil_actual else "",
                    placeholder="Ej: 2 Grúas autopropulsadas 50Tn, 3 Camiones pluma...",
                    height=120
                )

            with col_medios2:
                herramientas_software = st.text_area(
                    "Herramientas y Software",
                    value=perfil_actual['herramientas_software'] if perfil_actual else "",
                    placeholder="Ej: AutoCAD, PVSyst, MS Project...",
                    height=120
                )

            # Equipo profesional
            st.markdown("### 👥 Personal Técnico Asignado")

            # Cargar equipo existente o inicializar con campos ampliados
            equipo_actual = perfil_actual['equipo_tecnico'] if perfil_actual else []
            if not equipo_actual:
                equipo_actual = [{'nombre': '', 'cargo': '', 'titulacion': '', 'experiencia': 10, 'experiencia_detalle': '', 'certificaciones_personales': ''}] * 3

            num_tecnicos = st.number_input("Número de técnicos", 1, 20, len(equipo_actual))

            # Ajustar la lista si cambió el número
            while len(equipo_actual) < num_tecnicos:
                equipo_actual.append({'nombre': '', 'cargo': '', 'titulacion': '', 'experiencia': 10, 'experiencia_detalle': '', 'certificaciones_personales': ''})
            while len(equipo_actual) > num_tecnicos:
                equipo_actual.pop()

            equipo_tecnico = []

            # Mostrar formulario para cada técnico
            for i in range(num_tecnicos):
                with st.expander(f"👨‍💼 Personal Técnico {i+1}", expanded=i < 3):
                    col1, col2 = st.columns(2)

                    with col1:
                        nombre = st.text_input(f"Nombre completo",
                                            value=equipo_actual[i].get('nombre', '') if i < len(equipo_actual) else "",
                                            key=f"nombre_{i}",
                                            placeholder="Ej: Juan Pérez García")

                        cargo = st.text_input(f"Cargo/Función",
                                            value=equipo_actual[i].get('cargo', '') if i < len(equipo_actual) else "",
                                            key=f"cargo_{i}",
                                            placeholder="Ej: Jefe de Obra / Ingeniero Senior")

                        titulacion = st.text_input(f"Titulación específica",
                                                 value=equipo_actual[i].get('titulacion', '') if i < len(equipo_actual) else "",
                                                 key=f"tit_{i}",
                                                 placeholder="Ej: Ingeniero Industrial Col. 12345")

                    with col2:
                        experiencia_anos = st.number_input(f"Años de experiencia", 1, 40,
                                                         equipo_actual[i].get('experiencia', 10) if i < len(equipo_actual) else 10,
                                                         key=f"exp_{i}")

                        certificaciones_personales = st.text_input(f"Certificaciones personales",
                                                                  value=equipo_actual[i].get('certificaciones_personales', '') if i < len(equipo_actual) else "",
                                                                  key=f"cert_{i}",
                                                                  placeholder="Ej: PMP, Soldadura, BIM Manager...")

                    experiencia_detalle = st.text_area(f"Experiencia destacada y proyectos relevantes",
                                                     value=equipo_actual[i].get('experiencia_detalle', '') if i < len(equipo_actual) else "",
                                                     key=f"exp_det_{i}",
                                                     placeholder="Describe proyectos destacados, especialización técnica, trabajos previos relevantes...",
                                                     height=100)

                    equipo_tecnico.append({
                        "nombre": nombre,
                        "cargo": cargo,
                        "titulacion": titulacion,
                        "experiencia": experiencia_anos,
                        "experiencia_detalle": experiencia_detalle,
                        "certificaciones_personales": certificaciones_personales
                    })

            # Documentos Anexos
            st.markdown("### 📎 Documentos Anexos")
            st.info("💡 Los documentos que subas aquí se incluirán automáticamente como anexos en todas tus memorias técnicas")

            # Cargar documentos existentes
            documentos_actuales = perfil_actual.get('documentos_anexos', []) if perfil_actual else []

            # Mostrar documentos existentes
            if documentos_actuales:
                st.markdown("#### 📋 Documentos guardados:")
                cols_docs = st.columns(3)
                for i, doc in enumerate(documentos_actuales):
                    with cols_docs[i % 3]:
                        st.markdown(f"""
                        **{doc.get('nombre', 'Sin nombre')}**
                        📁 *{doc.get('categoria', 'General')}*
                        📅 {doc.get('fecha_subida', 'N/A')}
                        """)
                        if st.checkbox(f"Eliminar", key=f"delete_doc_{i}"):
                            documentos_actuales.remove(doc)
                            st.rerun()

            st.markdown("#### 📤 Subir nuevos documentos")

            # Categorías de documentos
            categorias_documentos = [
                "Certificados ISO (9001, 14001, 45001, etc.)",
                "Certificados de Buena Ejecución",
                "Currículums del Personal Técnico",
                "Certificados de Formación",
                "Licencias y Autorizaciones",
                "Referencias de Proyectos",
                "Certificados de Maquinaria/Equipos",
                "Pólizas de Seguros",
                "Otros Documentos Técnicos"
            ]

            col_cat, col_files = st.columns([1, 2])

            with col_cat:
                categoria_seleccionada = st.selectbox(
                    "Categoría del documento",
                    categorias_documentos
                )

            with col_files:
                archivos_subidos = st.file_uploader(
                    "Seleccionar archivos",
                    type=['pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png'],
                    accept_multiple_files=True,
                    help="Formatos aceptados: PDF, Word, JPG, PNG"
                )

            # Procesar archivos subidos
            documentos_anexos = list(documentos_actuales)  # Copia de los existentes

            if archivos_subidos:
                st.markdown("#### 📝 Archivos a subir:")
                for archivo in archivos_subidos:
                    st.write(f"📄 **{archivo.name}** - {categoria_seleccionada}")

            # Botón guardar
            if st.form_submit_button("💾 GUARDAR PERFIL DE EMPRESA", use_container_width=True):
                # Procesar logo si se subió uno nuevo
                logo_path = perfil_actual.get('logo_path', "") if perfil_actual else ""
                if logo_file:
                    nuevo_logo_path = guardar_logo_usuario(st.session_state.user_email, logo_file)
                    if nuevo_logo_path:
                        logo_path = nuevo_logo_path
                        st.success(f"✅ Logo guardado correctamente: {logo_file.name}")
                    else:
                        st.error("❌ Error guardando el logo")

                # Procesar documentos anexos si se subieron nuevos
                if archivos_subidos:
                    nuevos_documentos = guardar_documentos_anexos(
                        st.session_state.user_email,
                        archivos_subidos,
                        categoria_seleccionada
                    )
                    documentos_anexos.extend(nuevos_documentos)

                datos_perfil = {
                    'sector': sector,
                    'empleados': empleados,
                    'experiencia_anos': experiencia,
                    'certificaciones': certificaciones,
                    'otras_certificaciones': otras_certificaciones,
                    'experiencia_similar': experiencia_similar,
                    'logo_path': logo_path,
                    'medios_materiales': medios_materiales,
                    'herramientas_software': herramientas_software,
                    'equipo_tecnico': equipo_tecnico,
                    'documentos_anexos': documentos_anexos
                }

                if guardar_perfil_empresa(st.session_state.user_email, datos_perfil):
                    st.success("✅ Perfil de empresa guardado correctamente")
                    st.info("💡 Los datos se cargarán automáticamente en tus próximas memorias")
                    time.sleep(2)
                    st.rerun()
                else:
                    st.error("❌ Error al guardar el perfil")

    with tab2:
        st.markdown('<h2 class="section-header">Extracción Automática de Datos</h2>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            url_licitacion = st.text_input(
                "🔗 URL de la licitación (Plataforma de Contratación)",
                placeholder="https://contrataciondelestado.es/...",
                help="Pega el enlace directo de la licitación"
            )
        
        with col2:
            if st.button("🔍 Extraer Datos", type="primary"):
                if url_licitacion:
                    with st.spinner("Extrayendo datos de la plataforma..."):
                        datos = extraer_datos_licitacion(url_licitacion)
                        if datos:
                            st.session_state.datos_extraidos = datos
                            st.success("✅ Datos extraídos correctamente")
        
        if 'datos_extraidos' in st.session_state and st.session_state.datos_extraidos:
            st.info("📋 **Datos extraídos de la licitación:**")
            col3, col4 = st.columns(2)
            
            with col3:
                st.text_area("Objeto", value=st.session_state.datos_extraidos.get('objeto', ''), height=100)
                st.text_input("Expediente", value=st.session_state.datos_extraidos.get('expediente', ''))
                st.text_input("Organismo", value=st.session_state.datos_extraidos.get('organismo', ''))
            
            with col4:
                st.text_input("Presupuesto", value=st.session_state.datos_extraidos.get('presupuesto', ''))
                st.text_input("Plazo", value=st.session_state.datos_extraidos.get('plazo', ''))
    
    with tab2:
        st.markdown('<h2 class="section-header">Datos del Proyecto</h2>', unsafe_allow_html=True)

        # Cargar perfil de empresa automáticamente
        perfil_empresa = obtener_perfil_empresa(st.session_state.user_email)
        if perfil_empresa:
            st.success(f"✅ Datos de empresa cargados automáticamente desde tu perfil")
            if perfil_empresa['logo_path']:
                st.session_state['logo_path'] = perfil_empresa['logo_path']
        else:
            st.warning("⚠️ Completa primero tu perfil de empresa en la primera pestaña para cargar automáticamente los datos.")

        st.markdown('<h2 class="section-header">Datos de la Licitación</h2>', unsafe_allow_html=True)

        col3, col4 = st.columns(2)
        
        with col3:
            st.session_state.objeto = st.text_area("Objeto del Contrato *",
                                value=st.session_state.get('objeto', ''),
                                height=120,
                                placeholder="Ej: Instalación fotovoltaica de 500kW...")

            st.session_state.expediente = st.text_input("Número de expediente *",
                                      value=st.session_state.get('expediente', ''),
                                      placeholder="Ej: 2024/PA/001")

            st.session_state.organismo = st.text_input("Entidad Adjudicadora *",
                                     value=st.session_state.get('organismo', ''),
                                     placeholder="Ej: Ayuntamiento de...")

        with col4:
            st.session_state.presupuesto = st.text_input(
                "Presupuesto (formato español) *",
                value=st.session_state.get('presupuesto', ''),
                placeholder="Ej: 1.017.169,09",
                help="Formato español: puntos para miles, coma para decimales"
            )

            st.session_state.plazo = st.text_input("Plazo de Ejecución *",
                                value=st.session_state.get('plazo', ''),
                                placeholder="Ej: 12 meses")

            st.session_state.tipo_contrato = st.selectbox("Tipo de Contrato",
                                        ["Obras", "Servicios", "Suministros", "Mixto"],
                                        index=0)
    
    with tab3:
        st.markdown('<h2 class="section-header">Criterios de Juicio de Valor</h2>', unsafe_allow_html=True)
        
        st.warning("""
        ⚠️ **MUY IMPORTANTE**: Los criterios de valoración son el NÚCLEO de la memoria.
        El sistema desarrollará cada criterio en profundidad (3-5 páginas cada uno),
        relacionándolo directamente con el pliego técnico.
        """)
        
        num_criterios = st.number_input("Número de criterios de valoración", 1, 10, 3)
        
        st.session_state.criterios_valoracion = []
        total_puntos = 0
        
        for i in range(num_criterios):
            st.markdown(f"### Criterio {i+1}")
            col1, col2 = st.columns([3, 1])
            
            with col1:
                nombre_criterio = st.text_input(
                    f"Nombre del criterio {i+1}",
                    key=f"crit_nombre_{i}",
                    placeholder="Ej: Memoria técnica y metodología de trabajo"
                )
            
            with col2:
                puntos_criterio = st.number_input(
                    f"Puntos",
                    0, 100, 25,
                    key=f"crit_puntos_{i}"
                )
                total_puntos += puntos_criterio
            
            descripcion_criterio = st.text_area(
                f"Descripción/Aspectos a valorar",
                key=f"crit_desc_{i}",
                placeholder="Detalla qué aspectos se valorarán en este criterio",
                height=80
            )
            
            if nombre_criterio:
                st.session_state.criterios_valoracion.append({
                    'nombre': nombre_criterio,
                    'puntos': puntos_criterio,
                    'descripcion': descripcion_criterio
                })
        
        st.info(f"**Total puntos criterios técnicos:** {total_puntos}")
    
    with tab4:
        st.markdown('<h2 class="section-header">Documentos Técnicos</h2>', unsafe_allow_html=True)
        
        col5, col6 = st.columns(2)
        
        with col5:
            archivo_ppt = st.file_uploader("PPT - Pliego de Prescripciones Técnicas *", 
                                          type=['pdf', 'docx'],
                                          help="Fundamental para relacionar con criterios")
            if archivo_ppt:
                with st.spinner("🔍 Analizando PPT con extracción avanzada..."):
                    if archivo_ppt.type == "application/pdf":
                        analisis_ppt = extraer_texto_pdf_avanzado(archivo_ppt)
                        st.session_state.analisis_ppt = analisis_ppt
                        st.session_state.texto_ppt = analisis_ppt.get('texto_completo', '')

                        # Mostrar resumen del análisis
                        if 'resumen_analisis' in analisis_ppt:
                            with st.expander("📊 Resumen del Análisis del PPT", expanded=True):
                                st.text(analisis_ppt['resumen_analisis'])

                        # Mostrar criterios detectados automáticamente
                        criterios_detectados = analisis_ppt.get('criterios_valoracion', [])
                        if criterios_detectados:
                            st.success(f"✅ PPT analizado - {len(criterios_detectados)} criterios detectados automáticamente")

                            if st.button("🎯 Auto-completar criterios desde PPT"):
                                # Autocompletar criterios desde el análisis
                                st.session_state.criterios_valoracion = []
                                for criterio in criterios_detectados:
                                    st.session_state.criterios_valoracion.append({
                                        'nombre': criterio['nombre'],
                                        'puntos': criterio['puntos'] if criterio['puntos'] > 0 else 20
                                    })
                                st.success(f"✅ {len(criterios_detectados)} criterios añadidos automáticamente")
                                st.experimental_rerun()
                        else:
                            st.success("✅ PPT analizado - Análisis estructural completo")

                        # Mostrar tablas extraídas
                        tablas = analisis_ppt.get('tablas', [])
                        if tablas:
                            with st.expander(f"📋 {len(tablas)} tablas extraídas"):
                                for tabla in tablas[:3]:  # Mostrar máximo 3 tablas
                                    st.write(f"**Tabla {tabla['pagina']}.{tabla['indice']+1}** - Tipo: {tabla['tipo']} (Relevancia: {tabla['relevancia']}%)")
                                    if tabla['relevancia'] > 50:
                                        st.dataframe(pd.DataFrame(tabla['datos'], columns=tabla['headers']))
                    else:
                        st.session_state.texto_ppt = ""
                        st.warning("⚠️ Solo se soporta análisis avanzado para archivos PDF")
        
        with col6:
            archivo_pcap = st.file_uploader("PCAP - Pliego Administrativo (opcional)", 
                                           type=['pdf', 'docx'])
        
        st.markdown('<h2 class="section-header">Configuración de Generación</h2>', unsafe_allow_html=True)
        
        col9, col10 = st.columns(2)
        
        with col9:
            num_paginas = st.slider(
                "📄 Número de páginas de la memoria",
                min_value=1,
                max_value=120,
                value=60,
                step=1,
                help="Mayor extensión = mayor detalle en criterios"
            )
            
            extension_info = calcular_extension_contenido(num_paginas)
            st.info(f"""
            📊 **Configuración seleccionada:**
            - Páginas: {num_paginas}
            - Nivel de detalle: {extension_info['detalle']}
            - Desarrollo por criterio: {num_paginas // (len(st.session_state.criterios_valoracion) + 5) if len(st.session_state.criterios_valoracion) > 0 else 0} páginas aprox.
            """)
        
        with col10:
            incluir_graficos = st.checkbox("Incluir gráficos y diagramas", value=True)
            incluir_calculos = st.checkbox("Incluir cálculos técnicos", value=True)
            incluir_anexos = st.checkbox("Incluir anexos completos", value=True)
            st.session_state.incluir_cronograma = st.checkbox("📅 Generar cronograma de proyecto", value=st.session_state.get('incluir_cronograma', True), help="Cronograma Gantt adaptado al sector del proyecto")
            formato_profesional = st.checkbox("Formato profesional avanzado", value=True)

            # Vista previa del cronograma si está activado
            if st.session_state.incluir_cronograma and st.session_state.get('objeto') and st.session_state.get('plazo'):
                if st.button("👁️ Vista Previa del Cronograma"):
                    sector_detectado = detectar_sector_proyecto(
                        st.session_state.get('objeto', ''),
                        st.session_state.get('texto_ppt', '')
                    )
                    datos_temp = {
                        'objeto': st.session_state.get('objeto', ''),
                        'plazo': st.session_state.get('plazo', '12 meses')
                    }
                    fig, df = generar_cronograma_proyecto(datos_temp, sector_detectado)
                    st.plotly_chart(fig, use_container_width=True)

                    with st.expander("📊 Detalles del cronograma"):
                        st.dataframe(df[['Fase', 'Duración', 'Porcentaje']], use_container_width=True)
    
    with tab5:
        st.markdown('<h2 class="section-header">Generación de Memoria</h2>', unsafe_allow_html=True)
        
        # AVISO IMPORTANTE
        st.markdown("""
        <div class="warning-box">
            <h3>⚠️ AVISO IMPORTANTE</h3>
            <p><strong>Este sistema utiliza Inteligencia Artificial para generar memorias técnicas.</strong></p>
            <p>El sistema se centrará en:</p>
            <ul>
                <li>✓ Desarrollar cada criterio de valoración en profundidad</li>
                <li>✓ Relacionar cada criterio con el pliego técnico</li>
                <li>✓ Justificar con los recursos de la empresa</li>
                <li>✓ Crear contenido técnico profesional</li>
            </ul>
            <p>Los documentos DEBEN ser revisados por personal técnico cualificado.</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Validación - Necesitamos definir las variables aquí ya que no están en el scope de tab5
        # Obtener datos del usuario logueado desde la base de datos
        conn = sqlite3.connect('memoria_usuarios.db')
        cursor = conn.cursor()
        cursor.execute('SELECT empresa, cif FROM usuarios WHERE email = ?', (st.session_state.user_email,))
        datos_usuario = cursor.fetchone()
        conn.close()

        # Recuperar valores de los campos necesarios
        razon_social = datos_usuario[0] if datos_usuario else st.session_state.user_data.get('empresa', '')
        cif = datos_usuario[1] if datos_usuario else ''
        objeto = st.session_state.get('objeto', '')
        expediente = st.session_state.get('expediente', '')
        organismo = st.session_state.get('organismo', '')
        presupuesto = st.session_state.get('presupuesto', '')
        plazo = st.session_state.get('plazo', '')
        
        datos_completos = (
            razon_social and cif and objeto and 
            expediente and organismo and
            len(st.session_state.criterios_valoracion) > 0
        )
        
        if not datos_completos:
            st.warning("⚠️ Completa todos los campos obligatorios (*) y define al menos un criterio")
        
        # Resumen antes de generar
        if datos_completos:
            st.markdown("### 📋 Resumen de la memoria a generar:")
            col_res1, col_res2, col_res3, col_res4 = st.columns(4)
            
            with col_res1:
                st.metric("Páginas", num_paginas if 'num_paginas' in locals() else 60)
            with col_res2:
                st.metric("Criterios", len(st.session_state.criterios_valoracion))
            with col_res3:
                st.metric("Presupuesto", presupuesto if presupuesto else "N/A")
            with col_res4:
                st.metric("Plazo", plazo)
            
            # Mostrar criterios
            st.markdown("**🎯 Criterios de valoración a desarrollar:**")
            for criterio in st.session_state.criterios_valoracion:
                st.write(f"• {criterio['nombre']} ({criterio['puntos']} puntos)")
        
        # Deshabilitar en modo demo
        is_demo = st.session_state.get('demo_mode', False)
        button_disabled = not datos_completos or is_demo

        # DEBUG: Mostrar todos los valores para identificar el problema
        with st.expander("🔍 DEBUG - Ver datos actuales"):
            st.write(f"**razon_social:** '{razon_social}' ({'✅' if razon_social else '❌'})")
            st.write(f"**cif:** '{cif}' ({'✅' if cif else '❌'})")
            st.write(f"**objeto:** '{objeto}' ({'✅' if objeto else '❌'})")
            st.write(f"**expediente:** '{expediente}' ({'✅' if expediente else '❌'})")
            st.write(f"**organismo:** '{organismo}' ({'✅' if organismo else '❌'})")
            st.write(f"**criterios_valoracion:** {len(st.session_state.criterios_valoracion)} ({'✅' if len(st.session_state.criterios_valoracion) > 0 else '❌'})")
            st.write(f"**is_demo:** {is_demo}")
            st.write(f"**datos_completos:** {datos_completos}")
            st.write(f"**button_disabled:** {button_disabled}")

        # Mostrar qué falta para habilitar el botón
        if button_disabled and not is_demo:
            st.warning("⚠️ **Para generar la memoria, completa:**")
            faltantes = []
            if not razon_social: faltantes.append("• Razón Social (Perfil de Empresa)")
            if not cif: faltantes.append("• CIF (Perfil de Empresa)")
            if not objeto: faltantes.append("• Objeto del contrato (Datos del Proyecto)")
            if not expediente: faltantes.append("• Expediente (Datos del Proyecto)")
            if not organismo: faltantes.append("• Organismo (Datos del Proyecto)")
            if len(st.session_state.criterios_valoracion) == 0: faltantes.append("• Al menos 1 criterio de valoración")

            for falta in faltantes:
                st.write(falta)
        elif is_demo:
            st.info("🎭 **Modo Demo** - Para generar memorias, regístrate como cliente")

        if st.button("🚀 GENERAR MEMORIA TÉCNICA COMPLETA", disabled=button_disabled, type="primary"):
            with st.spinner("Generando memoria técnica profesional centrada en criterios..."):
                # Obtener configuración de extensión
                extension_config = calcular_extension_contenido(num_paginas if 'num_paginas' in locals() else 60)

                # Obtener datos del perfil de empresa guardado
                perfil_empresa = obtener_perfil_empresa(st.session_state.user_email)

                if not perfil_empresa:
                    st.error("❌ Debes completar tu perfil de empresa primero")
                    st.stop()

                # Obtener datos de usuario básicos
                conn = sqlite3.connect('memoria_usuarios.db')
                cursor = conn.cursor()
                cursor.execute('SELECT nombre, empresa, cif FROM usuarios WHERE email = ?',
                             (st.session_state.user_email,))
                datos_usuario = cursor.fetchone()
                conn.close()

                # Preparar datos de empresa con perfil guardado
                datos_empresa = {
                    'razon_social': datos_usuario[1] if datos_usuario else '',
                    'cif': datos_usuario[2] if datos_usuario else '',
                    'sector': perfil_empresa['sector'],
                    'empleados': perfil_empresa['empleados'],
                    'experiencia': perfil_empresa['experiencia_anos'],
                    'experiencia_similar': perfil_empresa['experiencia_similar'],
                    'certificaciones': perfil_empresa['certificaciones'],
                    'otras_certificaciones': perfil_empresa['otras_certificaciones'],
                    'medios_materiales': perfil_empresa['medios_materiales'],
                    'herramientas_software': perfil_empresa['herramientas_software'],
                    'equipo_tecnico': perfil_empresa['equipo_tecnico']
                }
                
                datos_proyecto = {
                    'titulo': objeto,
                    'objeto': objeto,
                    'expediente': expediente,
                    'organismo': organismo,
                    'presupuesto': presupuesto,
                    'plazo': plazo,
                    'tipo_contrato': st.session_state.get('tipo_contrato', 'Obras')
                }
                
                progress = st.progress(0)
                
                # 1. GENERAR DESARROLLO DE CRITERIOS (Lo más importante)
                with st.spinner("🎯 Desarrollando criterios de valoración en profundidad..."):
                    secciones_criterios = generar_memoria_por_criterios(
                        datos_proyecto,
                        st.session_state.criterios_valoracion,
                        st.session_state.texto_ppt,
                        datos_empresa
                    )
                    progress.progress(60)
                
                # 2. Generar otras secciones
                with st.spinner("📝 Generando secciones complementarias..."):
                    # Aquí podrías generar otras secciones necesarias
                    progress.progress(80)
                
                # Compilar todo
                datos_completos_doc = {
                    'empresa': datos_empresa,
                    'proyecto': datos_proyecto,
                    'criterios': st.session_state.criterios_valoracion,
                    'secciones_criterios': secciones_criterios,
                    'extension': num_paginas if 'num_paginas' in locals() else 60
                }
                
                # Crear documento Word
                with st.spinner("📄 Creando documento Word profesional..."):
                    doc = Document()

                    # Añadir numeración de páginas
                    agregar_numeracion_paginas(doc)

                    # Usar el logo del perfil de empresa
                    logo_a_usar = perfil_empresa.get('logo_path') if perfil_empresa else None
                    if not logo_a_usar:
                        logo_a_usar = st.session_state.get('logo_path')

                    print(f"DEBUG: Logo a usar: {logo_a_usar}")  # Debug
                    print(f"DEBUG: Datos empresa: {datos_empresa}")  # Debug

                    # Siempre añadir encabezado (con o sin logo)
                    agregar_encabezado_con_logo(doc, logo_a_usar, datos_empresa)

                    # Crear portada profesional
                    crear_portada_profesional(doc, datos_proyecto, datos_empresa, logo_a_usar)
                    
                    # Índice completo con subapartados y páginas
                    doc.add_heading('ÍNDICE', 1)

                    # Crear tabla de contenidos estructurada
                    # La portada es página 1, el índice página 2, contenido empieza en página 3
                    indice_items = [
                        ('1. PRESENTACIÓN DE LA EMPRESA', 3),
                        ('   1.1. Datos generales', 3),
                        ('   1.2. Experiencia y trayectoria', 4),
                        ('   1.3. Organización y estructura', 5),
                        ('   1.4. Certificaciones y acreditaciones', 6)
                    ]

                    pagina_actual = 7  # Después de la presentación de empresa (4 páginas)
                    for i, criterio in enumerate(st.session_state.criterios_valoracion, 2):
                        criterio_nombre = criterio['nombre'].upper()
                        indice_items.append((f'{i}. {criterio_nombre}', pagina_actual))
                        indice_items.append((f'   {i}.1. Análisis del criterio', pagina_actual))
                        indice_items.append((f'   {i}.2. Metodología propuesta', pagina_actual + 1))
                        indice_items.append((f'   {i}.3. Experiencia específica', pagina_actual + 2))
                        indice_items.append((f'   {i}.4. Recursos asignados', pagina_actual + 3))
                        indice_items.append((f'   {i}.5. Planificación y control', pagina_actual + 4))
                        pagina_actual += 6

                    # Añadir cronograma al índice si está habilitado
                    num_seccion = len(st.session_state.criterios_valoracion) + 2
                    if st.session_state.get('incluir_cronograma', True):
                        indice_items.append((f'{num_seccion}. CRONOGRAMA DE EJECUCIÓN', pagina_actual))
                        indice_items.append((f'   {num_seccion}.1. Planificación temporal', pagina_actual))
                        indice_items.append((f'   {num_seccion}.2. Fases del proyecto', pagina_actual + 1))
                        pagina_actual += 3
                        num_seccion += 1

                    # Añadir sección de anexos si hay documentos
                    if perfil_empresa and perfil_empresa.get('documentos_anexos'):
                        num_anexos = len(set(doc.get('categoria', 'General') for doc in perfil_empresa['documentos_anexos']))
                        indice_items.append((f'{num_seccion}. ANEXOS', pagina_actual))
                        for i, categoria in enumerate(set(doc.get('categoria', 'General') for doc in perfil_empresa['documentos_anexos']), 1):
                            categoria_corta = categoria.split('(')[0].strip()
                            indice_items.append((f'   Anexo {i} - {categoria_corta}', pagina_actual + i - 1))

                    # Añadir índice con páginas
                    for item, pagina in indice_items:
                        p = doc.add_paragraph()
                        run1 = p.add_run(item)
                        if not item.startswith('   '):
                            run1.bold = True
                        # Añadir puntos de relleno
                        espacios_necesarios = 70 - len(item)
                        puntos = '.' * max(5, espacios_necesarios // 2)
                        p.add_run(f' {puntos} ')
                        run_pagina = p.add_run(str(pagina))
                        run_pagina.bold = True

                    doc.add_page_break()
                    
                    # Desarrollar cada criterio
                    for criterio in st.session_state.criterios_valoracion:
                        doc.add_heading(criterio['nombre'].upper(), 1)
                        if criterio['nombre'] in secciones_criterios:
                            contenido = secciones_criterios[criterio['nombre']]
                            if contenido:
                                for parrafo in contenido.split('\n\n'):
                                    if parrafo.strip():
                                        doc.add_paragraph(parrafo)
                        doc.add_page_break()

                    # Añadir cronograma si está habilitado
                    if st.session_state.get('incluir_cronograma', True):
                        try:
                            sector_detectado = detectar_sector_proyecto(objeto, st.session_state.get('texto_ppt', ''))
                            fig, df_cronograma = generar_cronograma_proyecto(datos_proyecto, sector_detectado)

                            # Crear cronograma directamente en Word (más confiable)
                            cronograma_creado = crear_cronograma_tabla_word(doc, df_cronograma, datos_proyecto)

                            if cronograma_creado:
                                doc.add_page_break()
                            else:
                                # Fallback: cronograma básico de texto
                                doc.add_heading('CRONOGRAMA DE EJECUCIÓN', 1)
                                doc.add_paragraph(
                                    f"Cronograma de ejecución para el proyecto \"{objeto}\" "
                                    f"con plazo de {datos_proyecto.get('plazo', 'N/A')}."
                                )

                                # Añadir tabla básica con las fases
                                tabla_basic = doc.add_table(rows=1, cols=2)
                                tabla_basic.style = 'Table Grid'

                                hdr_cells = tabla_basic.rows[0].cells
                                hdr_cells[0].text = 'Fase del Proyecto'
                                hdr_cells[1].text = 'Porcentaje Estimado'

                                fases_basicas = [
                                    ("Planificación y diseño", "20%"),
                                    ("Desarrollo/Ejecución principal", "60%"),
                                    ("Supervisión y control", "15%"),
                                    ("Entrega y cierre", "5%")
                                ]

                                for fase, porcentaje in fases_basicas:
                                    row_cells = tabla_basic.add_row().cells
                                    row_cells[0].text = fase
                                    row_cells[1].text = porcentaje

                                doc.add_page_break()

                        except Exception as e:
                            print(f"Error generando cronograma en documento: {e}")
                            # Añadir cronograma básico de texto si hay error
                            doc.add_heading('CRONOGRAMA DE EJECUCIÓN', 1)
                            doc.add_paragraph(
                                f"Cronograma de ejecución para el proyecto \"{objeto}\" "
                                f"con plazo de {datos_proyecto.get('plazo', 'N/A')}."
                            )
                            doc.add_page_break()

                    # Añadir sección de anexos si hay documentos
                    if perfil_empresa and perfil_empresa.get('documentos_anexos'):
                        generar_seccion_anexos(doc, perfil_empresa['documentos_anexos'])

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    progress.progress(100)
                
                # Contar documentos anexos
                num_anexos = len(perfil_empresa.get('documentos_anexos', [])) if perfil_empresa else 0

                st.success(f"✅ Memoria técnica de {num_paginas if 'num_paginas' in locals() else 60} páginas generada correctamente")
                st.info("📌 Los criterios de valoración han sido desarrollados en profundidad y relacionados con el pliego técnico")

                if num_anexos > 0:
                    st.info(f"📎 Se han incluido {num_anexos} documentos como anexos en la memoria técnica")

                if st.session_state.get('incluir_cronograma', True):
                    st.info("📅 Se ha incluido cronograma Gantt profesional adaptado al sector del proyecto")
                
                # Botón de descarga
                st.download_button(
                    label=f"📥 DESCARGAR MEMORIA TÉCNICA ({num_paginas if 'num_paginas' in locals() else 60} páginas)",
                    data=buffer.getvalue(),
                    file_name=f"Memoria_Tecnica_{expediente}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                # Vista previa de criterios

                with st.expander("👁️ Vista previa - Desarrollo de Criterios"):
                    for nombre, contenido in secciones_criterios.items():
                        st.markdown(f"### {nombre}")
                        if contenido:
                            st.write(contenido[:1000] + "...")

        # Mostrar advertencia si está en modo demo
        if is_demo and datos_completos:
            st.info("ℹ️ En modo demo solo puedes visualizar la interfaz. Para generar memorias reales, regístrate como cliente.")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666;">
        <p><strong>🧠 MEMOR.IA - Sistema Profesional de Generación de Memorias Técnicas con IA</strong></p>
        <p>⚠️ Los documentos generados deben ser revisados por personal técnico cualificado</p>
        <p>© 2025 MEMOR.IA | Desarrollado por Grupo Oclem</p>
    </div>
    """, unsafe_allow_html=True)

# ============ MAIN ============
def main():
    # Inicializar base de datos
    init_database()

    # Inicializar estado de sesión
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'mostrar_registro' not in st.session_state:
        st.session_state.mostrar_registro = False
    if 'mostrar_recuperacion' not in st.session_state:
        st.session_state.mostrar_recuperacion = False
    if 'mostrar_login' not in st.session_state:
        st.session_state.mostrar_login = True

    # Lógica de navegación
    if not st.session_state.logged_in:
        if st.session_state.mostrar_registro:
            mostrar_registro()
        elif st.session_state.mostrar_recuperacion:
            mostrar_recuperacion()
        else:
            mostrar_login()
    else:
        # Botón de cerrar sesión en la barra lateral
        with st.sidebar:
            st.markdown("### 👤 Usuario Activo")
            st.info(f"""
            **Nombre:** {st.session_state.user_data['nombre']}  
            **Empresa:** {st.session_state.user_data['empresa']}  
            **Rol:** {st.session_state.user_data['rol']}
            """)
            
            if st.button("🚪 Cerrar Sesión", use_container_width=True):
                st.session_state.logged_in = False
                st.session_state.user_data = None
                st.session_state.user_email = None
                st.rerun()
            
            st.markdown("---")
            st.markdown("### 📚 Ayuda")
            st.markdown("""
            **Pasos clave:**
            1. 📄 Sube el pliego técnico
            2. 🎯 Define criterios de valoración
            3. 🏢 Completa datos empresa
            4. 🚀 Genera la memoria
            
            **IMPORTANTE:**
            El sistema desarrollará cada
            criterio en profundidad,
            relacionándolo con el pliego.
            """)
        
        # Verificar si es administrador
        if st.session_state.user_data['rol'] == 'admin':
            mostrar_aplicacion_admin()
        else:
            mostrar_aplicacion()

if __name__ == "__main__":
    main()
